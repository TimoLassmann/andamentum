"""Render scribe blocks to a python-docx Document.

The .docx is a derived artifact: this is a one-way render. We never
read .docx back into scribe. Templates are honoured if provided
via Document.template; otherwise python-docx's default styles are used.

Inline markdown formatting (**bold**, *italic*, `code`) inside paragraph
content is converted to styled runs via parser.inline_runs(). Anything
fancier (links, images-in-prose) falls back to plain text.

For figures: if the image file exists on disk we embed it; if not we
emit a placeholder paragraph (validate() will already have flagged the
missing file separately). Width honours `metadata.width_in` if set,
defaulting to 5.5 inches.

For tables: emitted as a real Word table with bold header row when
`metadata.header_row` is true. The caption is added as a styled
paragraph immediately after the table.
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any

from docx import Document as DocxDocument
from docx.shared import Inches

from .parser import find_ai_markers, inline_runs

if TYPE_CHECKING:  # pragma: no cover
    from .api import Document
    from docx.document import Document as _DocxDoc


_AI_MARKER_KEYWORD = "andamentum:contains-ai-markers"


def _stamp_ai_marker_keyword(out: "_DocxDoc") -> None:
    """Tag the docx's core properties with an AI-provenance keyword.

    Visible in Word's File→Info pane and recoverable by any docx-reading
    tool. Surfaces AI-assistance to readers (editors, integrity workflows)
    without forcing them to parse the body text. Uses the standard Dublin
    Core ``keywords`` field — no custom-properties XML needed.
    """
    existing = (out.core_properties.keywords or "").strip()
    if _AI_MARKER_KEYWORD in existing:
        return
    out.core_properties.keywords = (
        f"{existing}, {_AI_MARKER_KEYWORD}" if existing else _AI_MARKER_KEYWORD
    )


def _document_contains_ai_markers(doc: "Document") -> bool:
    for blk in doc.query(type="paragraph"):
        if find_ai_markers(blk.content):
            return True
    return False


def _emit_paragraph(out: "_DocxDoc", content: str) -> None:
    para = out.add_paragraph()
    for text, styles in inline_runs(content):
        run = para.add_run(text)
        if "bold" in styles:
            run.bold = True
        if "italic" in styles:
            run.italic = True
        if "code" in styles:
            run.font.name = "Courier New"


def _emit_table(out: "_DocxDoc", metadata: dict[str, Any]) -> None:
    rows = metadata.get("rows", [])
    if not rows:
        return
    table = out.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    header_row = bool(metadata.get("header_row", True))
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = cell_text
            if header_row and r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    caption = metadata.get("caption", "")
    if caption:
        p = out.add_paragraph(caption)
        if "Caption" in out.styles:
            p.style = "Caption"


def render_to_docx(doc: "Document", output_path: str) -> None:
    """Render `doc` to `output_path` as a .docx file."""
    if doc.template:
        out = DocxDocument(doc.template)
    else:
        out = DocxDocument()

    if _document_contains_ai_markers(doc):
        _stamp_ai_marker_keyword(out)

    for blk in doc.query():
        if blk.type == "heading":
            level = int(blk.metadata.get("level", 1))
            out.add_heading(blk.content, level=level)
        elif blk.type == "paragraph":
            _emit_paragraph(out, blk.content)
        elif blk.type == "figure":
            path = blk.metadata.get("path", "")
            caption = blk.metadata.get("caption", "")
            width_in = blk.metadata.get("width_in") or 5.5
            if path and Path(path).exists():
                out.add_picture(path, width=Inches(float(width_in)))
            else:
                out.add_paragraph(f"[Missing figure: {path}]")
            if caption:
                p = out.add_paragraph(caption)
                if "Caption" in out.styles:
                    p.style = "Caption"
        elif blk.type == "table":
            _emit_table(out, blk.metadata)
        else:  # pragma: no cover
            raise ValueError(f"Unknown block type for docx render: {blk.type!r}")

    out.save(output_path)
