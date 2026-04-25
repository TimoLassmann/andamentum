"""Tests for python-docx rendering."""

import pytest
from docx import Document as DocxDocument

from andamentum.scribe.api import Document, Figure, Heading, Paragraph, Table


def test_render_writes_docx_file(monkeypatch, tmp_path):
    monkeypatch.setenv("SCRIBE_DIR", str(tmp_path))
    doc = Document.create(title="P", database="t")
    doc.append(Heading("Intro", level=1))
    doc.append(Paragraph("Body text."))
    out = tmp_path / "out.docx"

    doc.render(str(out), format="docx")

    assert out.exists()
    docx = DocxDocument(str(out))
    paragraph_texts = [p.text for p in docx.paragraphs]
    assert "Intro" in paragraph_texts
    assert "Body text." in paragraph_texts


def test_render_heading_uses_heading_style(monkeypatch, tmp_path):
    monkeypatch.setenv("SCRIBE_DIR", str(tmp_path))
    doc = Document.create(title="P", database="t")
    doc.append(Heading("Methods", level=2))
    out = tmp_path / "out.docx"
    doc.render(str(out), format="docx")

    docx = DocxDocument(str(out))
    headings = [
        p
        for p in docx.paragraphs
        if p.style is not None
        and p.style.name is not None
        and p.style.name.startswith("Heading")
    ]
    assert any("Methods" in p.text for p in headings)
    assert any(p.style is not None and p.style.name == "Heading 2" for p in headings)


def test_render_inline_bold_emits_bold_run(monkeypatch, tmp_path):
    monkeypatch.setenv("SCRIBE_DIR", str(tmp_path))
    doc = Document.create(title="P", database="t")
    doc.append(Paragraph("normal **bold** text"))
    out = tmp_path / "out.docx"
    doc.render(str(out), format="docx")

    docx = DocxDocument(str(out))
    para = next(p for p in docx.paragraphs if "bold" in p.text)
    bold_runs = [r for r in para.runs if r.bold]
    assert any(r.text == "bold" for r in bold_runs)


def test_render_inline_italic_emits_italic_run(monkeypatch, tmp_path):
    monkeypatch.setenv("SCRIBE_DIR", str(tmp_path))
    doc = Document.create(title="P", database="t")
    doc.append(Paragraph("normal *em* text"))
    out = tmp_path / "out.docx"
    doc.render(str(out), format="docx")

    docx = DocxDocument(str(out))
    para = next(p for p in docx.paragraphs if "em" in p.text)
    italic_runs = [r for r in para.runs if r.italic]
    assert any(r.text == "em" for r in italic_runs)


_MINIMAL_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xcf\xc0"
    b"\x00\x00\x03\x01\x01\x00\xc9\xfe\x92\xef\x00\x00\x00\x00IEND\xaeB`\x82"
)


def test_render_figure_inserts_image_when_file_exists(monkeypatch, tmp_path):
    monkeypatch.setenv("SCRIBE_DIR", str(tmp_path))
    img = tmp_path / "f.png"
    img.write_bytes(_MINIMAL_PNG)
    doc = Document.create(title="P", database="t")
    doc.append(Figure(path=str(img), caption="cap", label="fig:c"))
    out = tmp_path / "out.docx"
    doc.render(str(out), format="docx")

    docx = DocxDocument(str(out))
    assert any("cap" in p.text for p in docx.paragraphs)


def test_render_figure_with_explicit_width(monkeypatch, tmp_path):
    monkeypatch.setenv("SCRIBE_DIR", str(tmp_path))
    img = tmp_path / "f.png"
    img.write_bytes(_MINIMAL_PNG)
    doc = Document.create(title="P", database="t")
    doc.append(Figure(path=str(img), caption="c", label="fig:x", width_in=4.5))
    out = tmp_path / "out.docx"
    doc.render(str(out), format="docx")  # must not raise


def test_render_table_emits_word_table(monkeypatch, tmp_path):
    monkeypatch.setenv("SCRIBE_DIR", str(tmp_path))
    doc = Document.create(title="P", database="t")
    doc.append(
        Table(
            rows=[["Col A", "Col B"], ["1", "2"], ["3", "4"]],
            header_row=True,
            caption="demo",
            label="tab:demo",
        )
    )
    out = tmp_path / "out.docx"
    doc.render(str(out), format="docx")

    docx = DocxDocument(str(out))
    assert len(docx.tables) == 1
    table = docx.tables[0]
    assert table.rows[0].cells[0].text == "Col A"
    assert table.rows[1].cells[1].text == "2"
    assert any("demo" in p.text for p in docx.paragraphs)


def test_render_unknown_format_raises(monkeypatch, tmp_path):
    monkeypatch.setenv("SCRIBE_DIR", str(tmp_path))
    doc = Document.create(title="P", database="t")
    with pytest.raises(ValueError, match="Unsupported format"):
        doc.render(str(tmp_path / "x"), format="pdf")
