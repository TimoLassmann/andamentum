"""Integration tests for precise, document-level comment anchoring.

Guards against the regression where every comment anchored to the
paragraph's last run, so multiple comments in one paragraph cascaded
into overlapping degenerate ranges (which Word silently ignored).

Each comment must now bracket its OWN target text; a target that can't
be located must fail loud (no degenerate placement); and a target that
spans a heading→body paragraph boundary must anchor across paragraphs.
"""

from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document as DocxDocument

from andamentum.whetstone.docx.finalization import finalize_reviewed_document
from andamentum.whetstone.models import DocumentPatch


def _range_text(doc_xml: str, cid: int) -> tuple[str, bool]:
    """Return (wrapped_text, overlaps_another_marker) for comment *cid*."""
    m = re.search(
        rf'<w:commentRangeStart w:id="{cid}"/>(.*?)<w:commentRangeEnd w:id="{cid}"/>',
        doc_xml,
        re.DOTALL,
    )
    inner = m.group(1) if m else ""
    text = "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", inner, re.DOTALL))
    overlaps = "commentReference" in inner or "commentRangeStart" in inner
    return text, overlaps


def _comment_ids(doc_xml: str) -> list[int]:
    return sorted(
        {int(x) for x in re.findall(r'<w:commentRangeStart w:id="(\d+)"', doc_xml)}
    )


def test_multiple_comments_one_paragraph_no_cascade(tmp_path: Path) -> None:
    """Three comments in one paragraph each wrap their own span — the old
    code piled them all on the last run with overlapping ranges."""
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = DocxDocument()
    doc.add_paragraph(
        "The methods were significantly robust and the gold standard was "
        "carefully built and the results were clearly described."
    )
    doc.save(str(src))

    patches = [
        DocumentPatch(
            patch_type="comment",
            text_pattern="significantly robust",
            comment_text="a",
            explanation="a",
        ),
        DocumentPatch(
            patch_type="comment",
            text_pattern="gold standard",
            comment_text="b",
            explanation="b",
        ),
        DocumentPatch(
            patch_type="comment",
            text_pattern="clearly described",
            comment_text="c",
            explanation="c",
        ),
    ]
    _, res = finalize_reviewed_document(
        original_file_path=src, patches=patches, output_path=out, author="t"
    )
    assert res.applied_patches == 3

    with zipfile.ZipFile(out) as z:
        doc_xml = z.read("word/document.xml").decode("utf-8")

    ids = _comment_ids(doc_xml)
    assert len(ids) == 3
    wrapped = {}
    for cid in ids:
        text, overlaps = _range_text(doc_xml, cid)
        assert text.strip(), f"comment {cid} wraps no text (degenerate)"
        assert not overlaps, (
            f"comment {cid} overlaps another comment's marker (cascade)"
        )
        wrapped[cid] = text
    # Each comment wraps the right target (allowing token-granular edges).
    joined = " ".join(wrapped.values())
    assert "significantly robust" in joined
    assert "gold standard" in joined
    assert "clearly described" in joined


def test_unmatched_target_fails_loud_not_placed(tmp_path: Path) -> None:
    """A target absent from the document must be reported as failed and NOT
    placed anywhere (no degenerate range)."""
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = DocxDocument()
    doc.add_paragraph("The methods were robust and the dataset was built.")
    doc.save(str(src))

    patches = [
        DocumentPatch(
            patch_type="comment",
            text_pattern="robust",
            comment_text="ok",
            explanation="ok",
        ),
        DocumentPatch(
            patch_type="comment",
            text_pattern="text that is absent",
            comment_text="ghost",
            explanation="ghost",
        ),
    ]
    _, res = finalize_reviewed_document(
        original_file_path=src, patches=patches, output_path=out, author="t"
    )
    assert res.applied_patches == 1
    assert len(res.failed_patches) == 1

    with zipfile.ZipFile(out) as z:
        doc_xml = z.read("word/document.xml").decode("utf-8")
    # Exactly one comment range placed (the matching one); the ghost is absent.
    assert len(_comment_ids(doc_xml)) == 1


def test_target_spanning_heading_and_body(tmp_path: Path) -> None:
    """A target that crosses a heading→body paragraph boundary anchors across
    the two paragraphs (commentRangeStart in one, End in the next)."""
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = DocxDocument()
    doc.add_paragraph("1 Introduction")
    doc.add_paragraph("Studying a biological research question is hard.")
    doc.save(str(src))

    # Target as a finding's quote would arrive from harvested markdown.
    patches = [
        DocumentPatch(
            patch_type="comment",
            text_pattern="## 1 Introduction\n\nStudying a biological research question",
            comment_text="section-level note",
            explanation="section-level note",
        )
    ]
    _, res = finalize_reviewed_document(
        original_file_path=src, patches=patches, output_path=out, author="t"
    )
    assert res.applied_patches == 1

    with zipfile.ZipFile(out) as z:
        doc_xml = z.read("word/document.xml").decode("utf-8")
    # The range start and end live in DIFFERENT paragraphs.
    start_para = doc_xml.split('<w:commentRangeStart w:id="1"/>')[0].count(
        "<w:p "
    ) + doc_xml.split('<w:commentRangeStart w:id="1"/>')[0].count("<w:p>")
    end_para = doc_xml.split('<w:commentRangeEnd w:id="1"/>')[0].count(
        "<w:p "
    ) + doc_xml.split('<w:commentRangeEnd w:id="1"/>')[0].count("<w:p>")
    assert end_para > start_para, "expected the range to span two paragraphs"


def test_comment_anchors_in_body_not_prepended_report(tmp_path: Path) -> None:
    """A comment whose target text is ALSO restated in the prepended review
    report must anchor to the manuscript body, not the report.

    The report is prepended before a page break; the body follows it. We
    assert the comment range starts AFTER the page break.
    """
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    doc = DocxDocument()
    doc.add_paragraph("The funding section is a TODO placeholder in the manuscript.")
    doc.save(str(src))

    patches = [
        DocumentPatch(
            patch_type="comment",
            text_pattern="funding section is a TODO",
            comment_text="Funding is a placeholder.",
            explanation="Funding is a placeholder.",
        )
    ]
    # The review report restates the same phrase — without the fix the
    # comment would anchor onto this report text (it appears first).
    _, res = finalize_reviewed_document(
        original_file_path=src,
        patches=patches,
        review_summary="The funding section is a TODO and must be completed.",
        issues_count=1,
        output_path=out,
        author="t",
    )
    assert res.applied_patches == 1

    with zipfile.ZipFile(out) as z:
        doc_xml = z.read("word/document.xml").decode("utf-8")
    assert _comment_ids(doc_xml) == [1]
    page_break_pos = doc_xml.find('<w:br w:type="page"/>')
    range_pos = doc_xml.find('<w:commentRangeStart w:id="1"/>')
    assert page_break_pos != -1 and range_pos != -1
    assert range_pos > page_break_pos, "comment anchored in the report, not the body"


def test_comment_text_not_duplicated(tmp_path: Path) -> None:
    """The comment body must not contain a duplicated 'Note:' re-append of
    its own text (the dedup fix)."""
    from andamentum.whetstone.renderers.docx import _to_document_patches
    from andamentum.whetstone.models import DocumentPatch as DP
    from andamentum.whetstone.schemas import Finding, Quote, ReviewResult

    finding = Finding(
        title="Passive voice",
        severity="minor",
        confidence="high",
        rationale="Detected passive-voice construction. Prefer active voice.",
        quotes=[Quote(section_id="s1", char_start=0, char_end=5, text="robust")],
        source="deterministic",
        category="style:passive",
    )
    result = ReviewResult(summary="", findings=[], deterministic_findings=[finding])
    patches = _to_document_patches(result, DP)
    comment_patches = [p for p in patches if p.patch_type == "comment"]
    assert len(comment_patches) == 1
    p = comment_patches[0]
    # explanation == comment_text → editor appends no "Note:" duplicate.
    assert p.explanation == p.comment_text
    assert p.comment_text.count("Detected passive-voice construction") == 1
