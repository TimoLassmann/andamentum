"""Tests for the tiered provenance watermark.

Three layers:
- Invisible metadata: always on regardless of visible_watermark.
- customXml provenance part (docx only): always on; survives core-property
  scrubbing.
- Visible banner: ON by default for review-report renderers, can be
  suppressed via visible_watermark=False (e.g. for --apply-patches
  output).
"""

from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document as DocxDocument

from andamentum.whetstone._watermark import (
    BANNER_TITLE,
    DISCLAIMER_SHORT,
    PROVENANCE_NS,
    metadata_markdown_comment,
    read_provenance_markers,
    stamp_docx_core_properties,
    write_provenance_customxml,
)
from andamentum.whetstone.renderers.html import render_html
from andamentum.whetstone.renderers.markdown import render_markdown
from andamentum.whetstone.schemas import ReviewResult


class TestMarkdownWatermark:
    def test_invisible_metadata_always_present(self) -> None:
        md = render_markdown(
            ReviewResult(summary="x"),
            model="openai:gpt-5.4-nano",
            visible_watermark=False,
        )
        assert "<!-- andamentum-whetstone" in md
        assert "ai-generated: true" in md
        assert "openai:gpt-5.4-nano" in md

    def test_visible_banner_on_by_default(self) -> None:
        md = render_markdown(ReviewResult(summary="x"))
        assert BANNER_TITLE in md
        assert DISCLAIMER_SHORT in md

    def test_visible_banner_off_when_flag_false(self) -> None:
        md = render_markdown(ReviewResult(summary="x"), visible_watermark=False)
        assert BANNER_TITLE not in md
        # But the invisible metadata still appears.
        assert "<!-- andamentum-whetstone" in md

    def test_clean_document_still_says_so(self) -> None:
        md = render_markdown(ReviewResult())
        assert "looks clean" in md.lower()


class TestHtmlWatermark:
    def test_visible_banner_on_by_default(self) -> None:
        html = render_html(ReviewResult(summary="x"))
        assert BANNER_TITLE in html

    def test_visible_banner_off_when_flag_false(self) -> None:
        html = render_html(ReviewResult(summary="x"), visible_watermark=False)
        assert BANNER_TITLE not in html

    def test_clean_document_still_says_so(self) -> None:
        html = render_html(ReviewResult())
        assert "looks clean" in html.lower()


class TestDocxMetadataStamp:
    def test_stamp_writes_contributor_keyword_and_description(
        self, tmp_path: Path
    ) -> None:
        path = tmp_path / "test.docx"
        DocxDocument().save(str(path))

        stamp_docx_core_properties(path, model="anthropic:claude-haiku-4-5")

        doc = DocxDocument(str(path))
        assert "andamentum-whetstone" in (doc.core_properties.author or "")
        assert "andamentum:ai-generated" in (doc.core_properties.keywords or "")
        assert "andamentum-whetstone" in (doc.core_properties.comments or "")
        assert "anthropic:claude-haiku-4-5" in (doc.core_properties.author or "")

    def test_stamp_is_idempotent(self, tmp_path: Path) -> None:
        path = tmp_path / "test.docx"
        DocxDocument().save(str(path))

        stamp_docx_core_properties(path, model="openai:gpt-5.4-nano")
        first_author = DocxDocument(str(path)).core_properties.author

        stamp_docx_core_properties(path, model="openai:gpt-5.4-nano")
        second_author = DocxDocument(str(path)).core_properties.author

        # Second call must not append a duplicate entry.
        assert first_author == second_author

    def test_stamp_preserves_existing_metadata(self, tmp_path: Path) -> None:
        path = tmp_path / "test.docx"
        doc = DocxDocument()
        doc.core_properties.author = "Dr Smith"
        doc.core_properties.keywords = "manuscript, draft"
        doc.save(str(path))

        stamp_docx_core_properties(path, model="ollama:llama3")

        doc = DocxDocument(str(path))
        assert "Dr Smith" in (doc.core_properties.author or "")
        assert "andamentum-whetstone" in (doc.core_properties.author or "")
        assert "manuscript, draft" in (doc.core_properties.keywords or "")
        assert "andamentum:ai-generated" in (doc.core_properties.keywords or "")

    def test_stamp_on_missing_file_does_not_raise(self, tmp_path: Path) -> None:
        # Best-effort: failures must not propagate
        stamp_docx_core_properties(tmp_path / "nonexistent.docx", model=None)


def test_markdown_metadata_comment_format() -> None:
    comment = metadata_markdown_comment(model="ollama:llama3")
    assert comment.startswith("<!--")
    assert comment.endswith("-->")
    assert "ai-generated: true" in comment
    assert "ollama:llama3" in comment


class TestCustomXmlProvenance:
    def test_customxml_part_is_written(self, tmp_path: Path) -> None:
        """The customXml provenance part is added to the docx zip."""
        path = tmp_path / "test.docx"
        DocxDocument().save(str(path))

        ok = write_provenance_customxml(path, model="openai:gpt-5.4-nano")
        assert ok is True

        with zipfile.ZipFile(path, "r") as zf:
            names = zf.namelist()
            assert "customXml/andamentum-provenance.xml" in names
            body = zf.read("customXml/andamentum-provenance.xml").decode("utf-8")
            assert PROVENANCE_NS in body
            assert "openai:gpt-5.4-nano" in body
            assert 'ai-generated="true"' in body
            # Content-types override declared.
            ct = zf.read("[Content_Types].xml").decode("utf-8")
            assert "/customXml/andamentum-provenance.xml" in ct

    def test_customxml_is_idempotent(self, tmp_path: Path) -> None:
        """Repeated writes produce one customXml part, refreshed in place."""
        path = tmp_path / "test.docx"
        DocxDocument().save(str(path))

        write_provenance_customxml(path, model="ollama:llama3")
        write_provenance_customxml(path, model="ollama:llama3")

        with zipfile.ZipFile(path, "r") as zf:
            prov_entries = [
                n for n in zf.namelist() if n == "customXml/andamentum-provenance.xml"
            ]
            assert len(prov_entries) == 1
            # And only one content-types override (no duplicate).
            ct = zf.read("[Content_Types].xml").decode("utf-8")
            assert ct.count("/customXml/andamentum-provenance.xml") == 1

    def test_customxml_survives_core_properties_scrub(self, tmp_path: Path) -> None:
        """Clearing core-properties keeps the customXml provenance intact."""
        path = tmp_path / "test.docx"
        DocxDocument().save(str(path))
        stamp_docx_core_properties(path, model="openai:gpt-5.4-nano")

        # Simulate "scrub metadata" — clear core-properties fields.
        doc = DocxDocument(str(path))
        doc.core_properties.author = ""
        doc.core_properties.keywords = ""
        doc.core_properties.comments = ""
        doc.save(str(path))

        markers = read_provenance_markers(path)
        assert markers["readable"] is True
        # Core-properties markers cleared by the scrub:
        assert markers["core_properties_marker"] is False
        assert markers["core_properties_author_marker"] is False
        # customXml layer survives:
        custom = markers["customxml_provenance"]
        assert isinstance(custom, dict)
        assert custom["generator"] == "andamentum-whetstone"
        assert custom["model"] == "openai:gpt-5.4-nano"

    def test_customxml_refuses_non_zip_file(self, tmp_path: Path) -> None:
        """Best-effort: non-docx input returns False, doesn't raise."""
        path = tmp_path / "not-a-docx.txt"
        path.write_text("hello")
        assert write_provenance_customxml(path, model=None) is False

    def test_customxml_refuses_missing_file(self, tmp_path: Path) -> None:
        """Missing file returns False, doesn't raise."""
        path = tmp_path / "nonexistent.docx"
        assert write_provenance_customxml(path, model=None) is False

    def test_xml_escapes_special_chars(self, tmp_path: Path) -> None:
        """Hostile model strings can't break the XML."""
        path = tmp_path / "test.docx"
        DocxDocument().save(str(path))
        # Model string with chars that need escaping
        write_provenance_customxml(path, model='evil"><script>x</script>')

        with zipfile.ZipFile(path, "r") as zf:
            body = zf.read("customXml/andamentum-provenance.xml").decode("utf-8")
        # The literal < > " from the input must be escaped, not raw.
        assert "<script>" not in body
        assert "&quot;" in body
        assert "&lt;script&gt;" in body


class TestReadProvenanceMarkers:
    def test_read_finds_all_three_signals(self, tmp_path: Path) -> None:
        """A fully-stamped docx surfaces all marker types."""
        path = tmp_path / "test.docx"
        DocxDocument().save(str(path))

        stamp_docx_core_properties(path, model="openai:gpt-5.4-nano")

        markers = read_provenance_markers(path)
        assert markers["readable"] is True
        assert markers["core_properties_marker"] is True
        assert markers["core_properties_author_marker"] is True
        custom = markers["customxml_provenance"]
        assert isinstance(custom, dict)
        assert custom["generator"] == "andamentum-whetstone"
        assert custom["model"] == "openai:gpt-5.4-nano"
        assert custom["ai-generated"] == "true"

    def test_read_on_clean_docx(self, tmp_path: Path) -> None:
        """An un-stamped docx reports no markers but is readable."""
        path = tmp_path / "clean.docx"
        DocxDocument().save(str(path))

        markers = read_provenance_markers(path)
        assert markers["readable"] is True
        assert markers["core_properties_marker"] is False
        assert markers["core_properties_author_marker"] is False
        assert markers["customxml_provenance"] is None

    def test_read_on_missing_file(self, tmp_path: Path) -> None:
        """Missing file: readable=False, no markers."""
        markers = read_provenance_markers(tmp_path / "nonexistent.docx")
        assert markers["readable"] is False
        assert markers["core_properties_marker"] is False
        assert markers["customxml_provenance"] is None

    def test_read_on_non_zip(self, tmp_path: Path) -> None:
        """Plain text file: readable=False, no markers."""
        path = tmp_path / "not-a-docx.txt"
        path.write_text("hello")
        markers = read_provenance_markers(path)
        assert markers["readable"] is False
