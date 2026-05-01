"""Tests for ``andamentum-whetstone --apply-patches``.

The patch-only CLI path is the analogue of the standalone
``apply-patches.py`` from the manuscript-tools plugin: take a .docx, take
a JSON list of ``DocumentPatch``, write a tracked-changes .docx. These
tests exercise the full end-to-end CLI surface against a real .docx
fixture (no mocks), plus the argparse-time validation rules.
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from andamentum.whetstone.cli import _build_parser, _validate_args, main


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _make_docx(path: Path, paragraphs: list[str]) -> None:
    """Write a minimal .docx with the given paragraphs to ``path``."""
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))


def _read_docx_xml_text(path: Path) -> str:
    """Concatenate ALL ``<w:t>`` text in a .docx body.

    Unlike ``Paragraph.text``, this surfaces text inside ``<w:ins>``
    (tracked insertions), which is exactly what the patch editor writes
    when applying a text_edit patch.
    """
    doc = Document(str(path))
    body = doc.element.body
    return "".join(t.text or "" for t in body.iter(f"{{{NS['w']}}}t"))


def _tracked_insertions(path: Path) -> list[str]:
    """Return the text content of every ``<w:ins>`` element in the document."""
    doc = Document(str(path))
    body = doc.element.body
    out: list[str] = []
    for ins in body.iter(f"{{{NS['w']}}}ins"):
        out.append("".join(t.text or "" for t in ins.iter(f"{{{NS['w']}}}t")))
    return out


def _tracked_deletions(path: Path) -> list[str]:
    """Return the text content of every ``<w:del>`` element in the document."""
    doc = Document(str(path))
    body = doc.element.body
    out: list[str] = []
    for d in body.iter(f"{{{NS['w']}}}del"):
        # Tracked deletions store text in <w:delText> not <w:t>
        out.append("".join(t.text or "" for t in d.iter(f"{{{NS['w']}}}delText")))
    return out


# ── argparse-time validation ────────────────────────────────────────────


def test_apply_patches_requires_docx_input(tmp_path: Path) -> None:
    """--apply-patches rejects a non-.docx INPUT."""
    md = tmp_path / "draft.md"
    md.write_text("# hi")
    patches = tmp_path / "patches.json"
    patches.write_text("[]")

    parser = _build_parser()
    args = parser.parse_args(
        [
            str(md),
            "--apply-patches",
            str(patches),
            "--out",
            str(tmp_path / "out.docx"),
        ]
    )
    with pytest.raises(SystemExit):
        _validate_args(args)


def test_apply_patches_requires_docx_out(tmp_path: Path) -> None:
    """--apply-patches rejects a non-.docx --out."""
    src = tmp_path / "draft.docx"
    _make_docx(src, ["hello"])
    patches = tmp_path / "patches.json"
    patches.write_text("[]")

    parser = _build_parser()
    args = parser.parse_args(
        [
            str(src),
            "--apply-patches",
            str(patches),
            "--out",
            str(tmp_path / "out.md"),
        ]
    )
    with pytest.raises(SystemExit):
        _validate_args(args)


def test_apply_patches_rejects_missing_patches_file(tmp_path: Path) -> None:
    """--apply-patches rejects a path that doesn't exist."""
    src = tmp_path / "draft.docx"
    _make_docx(src, ["hello"])

    parser = _build_parser()
    args = parser.parse_args(
        [
            str(src),
            "--apply-patches",
            str(tmp_path / "missing.json"),
            "--out",
            str(tmp_path / "out.docx"),
        ]
    )
    with pytest.raises(SystemExit):
        _validate_args(args)


def test_apply_patches_skips_model_requirement(tmp_path: Path) -> None:
    """In patch-only mode --model is NOT required."""
    src = tmp_path / "draft.docx"
    _make_docx(src, ["hello"])
    patches = tmp_path / "patches.json"
    patches.write_text("[]")

    parser = _build_parser()
    args = parser.parse_args(
        [
            str(src),
            "--apply-patches",
            str(patches),
            "--out",
            str(tmp_path / "out.docx"),
        ]
    )
    _validate_args(args)  # must not raise


def test_patch_report_without_apply_patches_rejected(tmp_path: Path) -> None:
    """--patch-report only makes sense in patch-only mode."""
    parser = _build_parser()
    args = parser.parse_args(
        [
            "paper.md",
            "--out",
            str(tmp_path / "out.md"),
            "--no-llm",
            "--patch-report",
            str(tmp_path / "report.md"),
        ]
    )
    with pytest.raises(SystemExit):
        _validate_args(args)


def test_patch_author_without_apply_patches_rejected(tmp_path: Path) -> None:
    """--patch-author only makes sense in patch-only mode."""
    parser = _build_parser()
    args = parser.parse_args(
        [
            "paper.md",
            "--out",
            str(tmp_path / "out.md"),
            "--no-llm",
            "--patch-author",
            "Alice",
        ]
    )
    with pytest.raises(SystemExit):
        _validate_args(args)


# ── end-to-end CLI run ─────────────────────────────────────────────────


def test_apply_patches_text_edit_end_to_end(tmp_path: Path) -> None:
    """A text_edit patch rewrites prose; the output .docx contains the new text."""
    src = tmp_path / "draft.docx"
    _make_docx(
        src,
        [
            "The quick brown fox jumps over the lazy dog.",
            "A second paragraph for context.",
        ],
    )

    patches = [
        {
            "patch_type": "text_edit",
            "text_pattern": "lazy dog",
            "new_text": "sleeping dog",
            "explanation": "More accurate description",
            "confidence": 0.9,
        }
    ]
    patches_path = tmp_path / "patches.json"
    patches_path.write_text(json.dumps(patches))

    out = tmp_path / "out.docx"
    main(
        [
            str(src),
            "--apply-patches",
            str(patches_path),
            "--out",
            str(out),
        ]
    )

    assert out.exists()
    # Track changes encode the diff token-by-token: the editor inserts
    # only the changed words ("sleeping") and deletes the original
    # ("lazy"), reusing the unchanged " dog" run. So check word-level
    # rather than full-phrase containment.
    inserted = " ".join(_tracked_insertions(out))
    deleted = " ".join(_tracked_deletions(out))
    assert "sleeping" in inserted
    assert "lazy" in deleted


def test_apply_patches_with_report_prepends_section(tmp_path: Path) -> None:
    """--patch-report prepends a review-summary section to the output."""
    src = tmp_path / "draft.docx"
    _make_docx(src, ["Some text to review."])

    patches_path = tmp_path / "patches.json"
    patches_path.write_text("[]")
    report = tmp_path / "report.md"
    report.write_text("Reviewer noticed nothing structural.")

    out = tmp_path / "out.docx"
    main(
        [
            str(src),
            "--apply-patches",
            str(patches_path),
            "--out",
            str(out),
            "--patch-report",
            str(report),
        ]
    )

    text = _read_docx_xml_text(out)
    assert "Reviewer noticed nothing structural." in text
    assert "Some text to review." in text


def test_apply_patches_invalid_json_array(tmp_path: Path) -> None:
    """Patches file must be a JSON array."""
    src = tmp_path / "draft.docx"
    _make_docx(src, ["hello"])
    patches_path = tmp_path / "patches.json"
    patches_path.write_text(json.dumps({"patches": []}))  # object, not array

    with pytest.raises(SystemExit) as excinfo:
        main(
            [
                str(src),
                "--apply-patches",
                str(patches_path),
                "--out",
                str(tmp_path / "out.docx"),
            ]
        )
    assert excinfo.value.code == 1


def test_apply_patches_invalid_patch_object(tmp_path: Path) -> None:
    """A malformed patch entry yields a config-error exit (1)."""
    src = tmp_path / "draft.docx"
    _make_docx(src, ["hello"])
    patches_path = tmp_path / "patches.json"
    # Missing required fields for text_edit (no new_text, no explanation)
    patches_path.write_text(
        json.dumps([{"patch_type": "text_edit", "text_pattern": "hello"}])
    )

    with pytest.raises(SystemExit) as excinfo:
        main(
            [
                str(src),
                "--apply-patches",
                str(patches_path),
                "--out",
                str(tmp_path / "out.docx"),
            ]
        )
    assert excinfo.value.code == 1
