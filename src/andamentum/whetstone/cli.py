"""CLI for whetstone v2 — ``andamentum-whetstone``.

Single positional argument (the input document, anything ``harvest`` can
read), one or more ``--out`` files (format inferred from the extension),
and a small set of well-named options. ``--help`` is exhaustive on
purpose: this is the surface most users will touch.

Exit codes:
    0 — success
    1 — configuration / argument error
    2 — input could not be loaded (harvest failure, file not found)
    3 — review pipeline failed
"""

from __future__ import annotations

import argparse
import asyncio
import logging
import os
import sys
import time
from pathlib import Path
from typing import NoReturn, Sequence

from rich.console import Console
from rich.logging import RichHandler
from rich.panel import Panel
from rich.table import Table

from andamentum.core import DEFAULT_EMBEDDING_MODEL

_LOGGER_NAME = "andamentum.whetstone"


_HELP_DESCRIPTION = """\
Sharpen YOUR OWN DRAFT with structured review.

Whetstone is NOT a peer-review tool. Do not run it on manuscripts,
grants, or any other documents shared with you in confidence (as a
journal reviewer, grant panel member, examiner, or editor). Most
publishers and funders currently prohibit sharing such documents
with AI tools, including cloud LLMs. See RESPONSIBLE_USE.md.

INPUT
    Path or URL to a document. Supported: .pdf, .docx, .html, .md, .txt
    (and any URL; HTML article-vs-listing is auto-detected).

OUTPUTS
    --out FILE         An output file. Repeat for multiple formats.
                       Format inferred from the extension:
                         .md   → markdown report
                         .html → HTML report (via typeset)
                         .docx → Word with track-changes + comments
                       For .docx: if INPUT is .docx we use it directly;
                       otherwise a clean .docx is generated from the
                       harvested text and patches are applied to that.

PATCH-ONLY MODE
    --apply-patches PATH       Skip the review pipeline entirely and apply
                               a pre-built JSON patch list to a .docx.
                               INPUT must be a local .docx and there must
                               be exactly one --out FILE ending in .docx.
                               PATH is a JSON file containing an array of
                               DocumentPatch objects (patch_type one of
                               text_edit | comment | document_analysis).
                               --model is not required.
    --patch-author NAME        Author name for the track-changes
                               attribution (default: Reviewer).
    --patch-report PATH        Optional markdown file prepended to the
                               output as a review summary section.

REVIEW OPTIONS
    --mode {review,panel,guidelines,custom}
                               Pipeline to run.
                                 review (default) — lens-based critical
                                   review with reflection loop.
                                 panel — simulate 3-5 fictional expert
                                   reviewers, scored, with a panel
                                   synthesis. Costs ~2N+2 LLM calls.
                                 guidelines — extract checkable items
                                   from a journal's author guidelines
                                   and evaluate each against the
                                   manuscript. Requires --guidelines.
                                 custom — evaluate the manuscript
                                   against caller-supplied criteria
                                   (one LLM call). Requires --criteria.
    --model MODEL              pydantic-ai model id, e.g.
                               openai:gpt-5.4-nano,
                               ollama:gemma4:31b-nvfp4.
                               Required unless --no-llm is given.
    --editor                   Also generate concrete edits (rewrites).
                               Default off — adds one LLM call per section.
                               Ignored in panel mode.
    --editor-criteria LIST     Comma-separated. Default:
                               clarity,concision,grammar
    --no-challenge             Skip the refutation phase
                               (faster, slightly less reliable findings).
                               Ignored in panel mode.
    --perspectives LIST        Comma-separated lens names.
                               Default: rigorous
                               Available: rigorous, writer, methodology,
                                          statistician, consistency,
                                          claim_evidence, overclaim, strunk
                               Notes:
                                 consistency reads the WHOLE document
                                   (cross-section drift / contradictions)
                                 strunk applies Elements of Style rules
                                   via a per-rule pydantic-graph sub-graph
                                 claim_evidence runs only on Abstract /
                                   Results / Discussion / Conclusion
                                 overclaim flags unsupported strength
                                   language ("first / novel / dramatic")
                               Ignored in panel mode.
    --rounds N                 Hard cap on rounds of the reflection–
                               investigation loop. Default: 3.
                               The loop typically exits earlier when the
                               senior reviewer says "nothing more to do".
                               Ignored in panel mode.
    --n-experts N              In panel mode, how many experts to
                               generate. Default 4.
    --panel-disciplines LIST   In panel mode, an explicit comma-separated
                               list of disciplines (skips keyword
                               extraction).
    --guidelines TEXT          In guidelines mode, the journal author
                               guidelines as text. Use ``@path/to/file``
                               to read from a file. Required when
                               ``--mode guidelines``.
    --criteria LIST            In custom mode, semicolon-separated list
                               of criteria (e.g. "originality; depth of
                               literature; clarity of methods"). May be
                               repeated to supply criteria one at a time.
                               Required when ``--mode custom``.
    --no-llm                   Run only the deterministic structural pass
                               (citations, terms, numerics, cross-refs).
                               No --model required. Free, instant.
    -v, --verbose              Print phase-by-phase progress to stderr.
"""


_HELP_EXAMPLES = """\
EXAMPLES
    # Quick deterministic-only check (no LLM, no cost)
    andamentum-whetstone paper.pdf --no-llm --out review.md

    # Apply a pre-built JSON patch list to a .docx (no LLM)
    andamentum-whetstone draft.docx \\
        --apply-patches patches.json --out draft.reviewed.docx \\
        --patch-author "Claude" --patch-report review.md

    # Full critical review of a PDF, output in three formats
    andamentum-whetstone paper.pdf \\
        --model openai:gpt-5.4-nano \\
        --out review.md --out review.html --out paper.reviewed.docx

    # Edit-mode: get concrete rewrites in a tracked-change Word doc
    andamentum-whetstone draft.docx \\
        --model openai:gpt-5.4-nano --editor \\
        --out draft.reviewed.docx

    # Multi-lens panel on an arXiv paper
    andamentum-whetstone https://arxiv.org/pdf/1901.01753 \\
        --model openai:gpt-5.4-nano \\
        --perspectives rigorous,statistician,writer \\
        --out panel-review.html

    # Local model, deeper reflection loop (4 rounds), all outputs
    andamentum-whetstone manuscript.pdf \\
        --model ollama:gemma4:31b-nvfp4 --rounds 4 --editor \\
        --out review.md --out review.html --out manuscript.reviewed.docx

EXIT CODES
    0  success
    1  configuration / argument error
    2  input could not be loaded (harvest failure, file not found)
    3  review pipeline failed
"""


_SUPPORTED_OUT_EXTENSIONS = {".md", ".html", ".docx"}


def _build_parser() -> argparse.ArgumentParser:
    """Build the argparse parser with the full help text."""
    parser = argparse.ArgumentParser(
        prog="andamentum-whetstone",
        description=_HELP_DESCRIPTION,
        epilog=_HELP_EXAMPLES,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument(
        "input",
        help="Path or URL to the document to review.",
    )
    parser.add_argument(
        "--out",
        action="append",
        type=Path,
        required=True,
        metavar="FILE",
        help="Output file. Repeat for multiple formats. "
        "Format inferred from extension (.md / .html / .docx).",
    )
    parser.add_argument(
        "--mode",
        choices=("review", "panel", "guidelines", "custom"),
        default="review",
        help=(
            "Pipeline to run. review (default): lens-based critical "
            "review. panel: simulate 3-5 fictional expert reviewers. "
            "guidelines: evaluate document against journal author "
            "guidelines (requires --guidelines). custom: evaluate "
            "against caller-supplied criteria (requires --criteria)."
        ),
    )
    parser.add_argument(
        "--model",
        metavar="MODEL",
        help="pydantic-ai model id (e.g. openai:gpt-5.4-nano, "
        "ollama:gemma4:31b-nvfp4). Required unless --no-llm.",
    )
    parser.add_argument(
        "--v3",
        action="store_true",
        help=(
            "Use the experimental v3 whole-document pipeline (claims digest + "
            "SPECS criteria + gap loop) instead of the v2 section-by-section "
            "review. Requires --model. Most v2-only flags (perspectives, panel, "
            "editor, …) are ignored; --document-type selects the criterion set "
            "and --rounds sets the gap-loop cap."
        ),
    )
    parser.add_argument(
        "--embedding-model",
        default=DEFAULT_EMBEDDING_MODEL,
        metavar="MODEL",
        help=(
            "Local Ollama embedding model used by the Consolidate phase to "
            "spot similar comments. Default: "
            f"{DEFAULT_EMBEDDING_MODEL}. Ollama must be running."
        ),
    )
    parser.add_argument(
        "--n-experts",
        type=int,
        default=4,
        metavar="N",
        help="In panel mode, how many experts to generate. Default: 4.",
    )
    parser.add_argument(
        "--panel-disciplines",
        default="",
        metavar="LIST",
        help=(
            "In panel mode, an explicit comma-separated list of "
            "disciplines (skips keyword extraction)."
        ),
    )
    parser.add_argument(
        "--i-am-the-author",
        action="store_true",
        help=(
            "Required for --mode panel. Affirms that the document being "
            "reviewed is your own draft, not a manuscript shared with you "
            "confidentially (as a peer reviewer, examiner, or grant panel "
            "member). Set ANDAMENTUM_PANEL_OWN_AUTHOR=1 to pre-affirm."
        ),
    )
    parser.add_argument(
        "--confirm-own-draft",
        action="store_true",
        help=(
            "Bypass the confidentiality-marker tripwire (which refuses "
            "to run when the document contains phrases like 'Manuscript "
            "ID:', 'Reviewer Instructions', 'Editorial Office', etc.). "
            "Pass this flag only when the matched marker is a legitimate "
            "false positive in your own draft."
        ),
    )
    parser.add_argument(
        "--document-type",
        choices=(
            "auto",
            "academic",
            "external_communication",
            "essay",
            "tutorial",
            "creative",
            "general",
        ),
        default="auto",
        help=(
            "What kind of document this is. 'auto' (default) runs a "
            "one-shot classifier using --model. Explicit values skip the "
            "classifier. Six document types route to different criterion "
            "sets (academic / external_communication / essay / tutorial "
            "/ creative / general); the journal-specific checklist (CoI "
            "/ data / ethics / abstract / keywords / H1 title) fires "
            "only for 'academic'; synthesis vocabulary adapts to type."
        ),
    )
    parser.add_argument(
        "--guidelines",
        default="",
        metavar="TEXT",
        help=(
            "In guidelines mode, the journal author guidelines text. "
            "Use ``@path/to/file`` to read from a file. Required when "
            "--mode guidelines."
        ),
    )
    parser.add_argument(
        "--criteria",
        action="append",
        default=None,
        metavar="LIST",
        help=(
            "In custom mode, criteria to evaluate against. Either a "
            "single semicolon-separated string or repeat the flag to "
            "supply criteria one at a time. Required when --mode custom."
        ),
    )
    parser.add_argument(
        "--editor",
        action="store_true",
        help="Also generate concrete edits (rewrites). Adds one LLM call per section.",
    )
    parser.add_argument(
        "--editor-criteria",
        default="clarity,concision,grammar",
        metavar="LIST",
        help="Comma-separated editor criteria. Default: clarity,concision,grammar",
    )
    parser.add_argument(
        "--no-challenge",
        action="store_true",
        help="Skip the refutation phase (faster, slightly less reliable findings).",
    )
    parser.add_argument(
        "--perspectives",
        default="rigorous",
        metavar="LIST",
        help=(
            "Comma-separated lens names. "
            "Available: rigorous, writer, methodology, statistician, "
            "consistency, claim_evidence, overclaim, strunk. "
            "Default: rigorous"
        ),
    )
    parser.add_argument(
        "--rounds",
        type=int,
        default=3,
        metavar="N",
        help=(
            "Hard cap on rounds of the reflection-investigation loop. "
            "Default: 3. The loop typically exits earlier."
        ),
    )
    parser.add_argument(
        "--no-llm",
        action="store_true",
        help="Run only the deterministic structural pass. No --model required.",
    )
    parser.add_argument(
        "--no-proofread",
        action="store_true",
        help=(
            "Disable the deterministic proofread pass (weasel words, passive "
            "voice, duplicate words, weak sentence openers). On by default — "
            "its findings flow into the same report as the LLM-driven lenses."
        ),
    )
    parser.add_argument(
        "--check-novelty",
        action="store_true",
        help=(
            "Verify the manuscript's novelty claims against the literature "
            "via deep_research. Adds 1 + (3-5)*2 LLM calls + several web "
            "fetches. Disabled by default."
        ),
    )
    parser.add_argument(
        "--novelty-search-depth",
        type=int,
        default=2,
        metavar="N",
        help=(
            "Search depth for --check-novelty. 1=quick, 2=balanced, 3=thorough. "
            "Default: 2."
        ),
    )
    parser.add_argument(
        "--persist-novelty-cache",
        action="store_true",
        help=(
            "Persist per-claim novelty-check results to "
            "~/.cache/whetstone/novelty/ so re-runs on the same draft are "
            "cheap. Off by default — hashed digests of unpublished novelty "
            "claims should not sit on disk unless you explicitly opt in."
        ),
    )
    visible_wm = parser.add_mutually_exclusive_group()
    visible_wm.add_argument(
        "--visible-watermark",
        dest="visible_watermark",
        action="store_true",
        default=True,
        help=(
            "Include a visible 'AI-generated review content' banner in the "
            "review report (default for review-mode output)."
        ),
    )
    visible_wm.add_argument(
        "--no-visible-watermark",
        dest="visible_watermark",
        action="store_false",
        help=(
            "Suppress the visible banner. Invisible provenance metadata "
            "(docx core properties, HTML <meta>, markdown HTML-comment) "
            "is still written regardless of this flag."
        ),
    )
    parser.add_argument(
        "--apply-patches",
        type=Path,
        default=None,
        metavar="PATH",
        help=(
            "Skip the review pipeline; apply a pre-built JSON patch list "
            "to INPUT (which must be a local .docx). PATH points to a "
            "JSON array of DocumentPatch objects."
        ),
    )
    parser.add_argument(
        "--patch-author",
        default=None,
        metavar="NAME",
        help=(
            "Author name for track-changes when --apply-patches is used. "
            "Default: 'andamentum-whetstone (AI)'. Overriding to a custom "
            "name requires --allow-author-override (misrepresenting "
            "AI-generated edits as a human reviewer's may constitute "
            "research misconduct)."
        ),
    )
    parser.add_argument(
        "--allow-author-override",
        action="store_true",
        help=(
            "Explicitly authorise a non-default --patch-author value. "
            "Required when --patch-author is set to anything other than "
            "the AI-attribution default."
        ),
    )
    parser.add_argument(
        "--patch-report",
        type=Path,
        default=None,
        metavar="PATH",
        help=(
            "Optional markdown file prepended as a review-summary section "
            "when --apply-patches is used."
        ),
    )
    parser.add_argument(
        "-v",
        "--verbose",
        action="store_true",
        help="Print progress to stderr.",
    )
    return parser


def _validate_args(args: argparse.Namespace) -> None:
    """Argparse-time validation that argparse can't express directly.

    Raises ``SystemExit(1)`` with a friendly message on any violation.
    """
    if args.apply_patches is not None:
        _validate_apply_patches_args(args)
        return

    if args.patch_report is not None:
        _die(1, "--patch-report is only valid with --apply-patches.")
    if args.patch_author is not None:
        _die(1, "--patch-author is only valid with --apply-patches.")
    if args.allow_author_override:
        _die(1, "--allow-author-override is only valid with --apply-patches.")

    if not args.model:
        _die(
            1,
            "--model is required.\nExamples:\n"
            "  --model openai:gpt-5.4-nano\n"
            "  --model ollama:gemma4:31b-nvfp4",
        )
    if args.mode == "panel" and not (
        args.i_am_the_author or os.environ.get("ANDAMENTUM_PANEL_OWN_AUTHOR") == "1"
    ):
        _die(
            1,
            "panel subcommand requires --i-am-the-author. Panel-mode output "
            "is shaped exactly like a journal peer-review report (3-5 "
            "fictional reviewer biosketches + Accept/Reject recommendation). "
            "Affirm that the document being reviewed is your own draft, not "
            "a manuscript shared with you confidentially. Pass "
            "--i-am-the-author, or set ANDAMENTUM_PANEL_OWN_AUTHOR=1.",
        )
    # --criteria and --guidelines are mutually exclusive — both route to
    # the unified criterion-input surface (Phase C); the API itself
    # enforces this with a ValueError, but we surface a friendlier
    # message earlier here.
    if args.criteria and args.guidelines:
        _die(
            1,
            "--criteria and --guidelines are mutually exclusive. Pass one "
            "or the other — both are routes to the same active-criteria slot.",
        )
    if args.rounds < 1:
        _die(1, "--rounds must be at least 1.")
    if args.n_experts < 1:
        _die(1, "--n-experts must be at least 1.")
    for out in args.out:
        ext = out.suffix.lower()
        if ext not in _SUPPORTED_OUT_EXTENSIONS:
            _die(
                1,
                f"Unsupported output extension {ext!r} for {out}. "
                f"Supported: {', '.join(sorted(_SUPPORTED_OUT_EXTENSIONS))}",
            )


def _validate_apply_patches_args(args: argparse.Namespace) -> None:
    """Validate the patch-only path. ``--apply-patches`` is set."""
    input_path = Path(args.input)
    if _is_url(args.input) or not input_path.exists():
        _die(
            1,
            "--apply-patches requires INPUT to be a local .docx file "
            f"(got {args.input!r}).",
        )
    if input_path.suffix.lower() != ".docx":
        _die(
            1,
            "--apply-patches requires INPUT to be a .docx file "
            f"(got {input_path.suffix!r}).",
        )
    if len(args.out) != 1 or args.out[0].suffix.lower() != ".docx":
        _die(
            1,
            "--apply-patches requires exactly one --out FILE ending in .docx.",
        )
    if not args.apply_patches.exists():
        _die(1, f"patches file not found: {args.apply_patches}")
    if args.patch_report is not None and not args.patch_report.exists():
        _die(1, f"patch report file not found: {args.patch_report}")
    if args.patch_author is not None and not args.allow_author_override:
        _die(
            1,
            "--patch-author overrides the AI-attribution default — "
            "pass --allow-author-override to confirm you intend this. "
            "Attributing AI-generated edits to a human name may "
            "constitute research misconduct under most institutional codes.",
        )
    if args.allow_author_override and args.patch_author is None:
        _die(1, "--allow-author-override requires --patch-author NAME.")


def _die(code: int, message: str) -> NoReturn:
    """Print to stderr and exit with the given code."""
    print(f"error: {message}", file=sys.stderr)
    sys.exit(code)


def _resolve_guidelines(value: str) -> str:
    """Resolve --guidelines value, expanding ``@path`` to file contents.

    A leading ``@`` selects file mode; everything else is treated as a
    literal guidelines string.

    Raises
    ------
    SystemExit
        With exit code 2 if the file cannot be read.
    """
    if value.startswith("@"):
        path = Path(value[1:])
        try:
            return path.read_text(encoding="utf-8")
        except OSError as exc:
            _die(2, f"could not read guidelines file {path}: {exc}")
    return value


def _parse_criteria(values: list[str]) -> list[str]:
    """Parse ``--criteria`` into a flat list of criterion strings.

    Each ``--criteria`` invocation may be a semicolon-separated string;
    multiple invocations stack. Empty / whitespace-only entries are
    silently dropped.
    """
    out: list[str] = []
    for raw in values:
        for chunk in raw.split(";"):
            cleaned = chunk.strip()
            if cleaned:
                out.append(cleaned)
    return out


async def _run(args: argparse.Namespace, console: Console) -> None:
    """Execute the review and write all requested outputs."""
    from .renderers import render_docx, render_html, render_markdown
    from .v3 import Criterion, review_document_v3
    from .v3.panel import run_panel_v3

    editor_criteria = tuple(
        c.strip() for c in args.editor_criteria.split(",") if c.strip()
    )
    panel_disciplines = tuple(
        d.strip() for d in args.panel_disciplines.split(",") if d.strip()
    )
    guidelines_text = _resolve_guidelines(args.guidelines) if args.guidelines else ""

    # Always-on one-line scope banner — short, unmissable, not gated on
    # --verbose. Goes to stderr so it doesn't pollute piped output.
    print(
        "Note: whetstone is for your own drafts. Not a peer-review tool. "
        "See RESPONSIBLE_USE.md.",
        file=sys.stderr,
    )

    logger = logging.getLogger(_LOGGER_NAME)
    if args.verbose:
        _log_run_config(console, args, (), editor_criteria)

    try:
        _t0 = time.perf_counter()
        if args.mode == "panel":
            # Panel mode runs its own graph (different shape — multi-
            # expert fan-out + synthesis); criterion-cascade kwargs do
            # not apply.
            from andamentum.harvest import extract as harvest_extract
            from pathlib import Path as _Path

            md = (
                await harvest_extract(_Path(args.input))
                if _Path(args.input).exists()
                else args.input
            )
            result = await run_panel_v3(
                md,
                model=args.model,
                n_experts=args.n_experts,
                panel_disciplines=list(panel_disciplines) or None,
            )
        else:
            # Default review path — Phase C unified criteria/guidelines
            # input is mutually exclusive (CLI validator enforces).
            v3_criteria: list[Criterion] | None = None
            v3_guidelines: str | None = None
            if args.criteria:
                parsed = _parse_criteria(args.criteria)
                v3_criteria = [Criterion(name=s, questions=[s]) for s in parsed]
            if args.guidelines:
                v3_guidelines = guidelines_text

            result = await review_document_v3(
                args.input,
                model=args.model,
                cap=args.rounds,
                document_type=args.document_type,
                confirm_own_draft=args.confirm_own_draft,
                criteria=v3_criteria,
                guidelines_text=v3_guidelines,
                editor=args.editor,
                editor_criteria=list(editor_criteria) if editor_criteria else None,
                check_novelty=args.check_novelty,
                novelty_search_depth=args.novelty_search_depth,
            )
        result.metrics.wall_seconds = time.perf_counter() - _t0
    except Exception as exc:
        # Confidentiality tripwire — its own error class so we can give a
        # focused error code (1: configuration) instead of pipeline-crash (3).
        from ._confidentiality import ConfidentialityMarkerError

        if isinstance(exc, ConfidentialityMarkerError):
            _die(1, str(exc))
        # Distinguish "couldn't load the input" from "review crashed"
        # so users know whether to fix their input or their config.
        msg = str(exc)
        if any(
            t in msg.lower() for t in ("fetch", "harvest", "no such file", "not found")
        ):
            _die(2, f"could not load input {args.input!r}: {exc}")
        else:
            _die(3, f"review pipeline failed: {exc}")

    # If any output is .docx, prepare a baseline .docx ONCE upfront
    # (rather than per-output) so we don't re-harvest if the user asked
    # for multiple .docx outputs.
    baseline_docx: Path | None = None
    baseline_is_temp = False
    if any(o.suffix.lower() == ".docx" for o in args.out):
        baseline_docx, baseline_is_temp = await _ensure_docx_source(args.input)

    written: list[Path] = []
    try:
        # ── Write each requested output ──────────────────────────────
        for out in args.out:
            ext = out.suffix.lower()
            out.parent.mkdir(parents=True, exist_ok=True)
            model_id = None if args.no_llm else args.model
            if ext == ".md":
                render_markdown(
                    result,
                    out,
                    model=model_id,
                    visible_watermark=args.visible_watermark,
                )
            elif ext == ".html":
                render_html(
                    result,
                    out,
                    model=model_id,
                    visible_watermark=args.visible_watermark,
                )
            elif ext == ".docx":
                assert baseline_docx is not None  # validated above
                render_docx(
                    result,
                    source_docx_path=baseline_docx,
                    output_path=out,
                    model=model_id,
                    visible_watermark=args.visible_watermark,
                )
            written.append(out)
            logger.info("[output] wrote %s", out)
    finally:
        # Clean up the auto-generated baseline docx (if we created one).
        if baseline_is_temp and baseline_docx is not None:
            try:
                baseline_docx.unlink()
            except OSError:
                pass

    _print_summary(console, result, written, v3=args.v3)

    # Disclosure reminder — always shown at end of run. AI assistance must
    # be disclosed in submitted artifacts per most journal / funder rules.
    from ._watermark import DISCLOSURE_REMINDER

    print(f"\nNote: {DISCLOSURE_REMINDER}", file=sys.stderr)


def _log_run_config(
    console: Console,
    args: argparse.Namespace,
    perspectives: tuple[str, ...],
    editor_criteria: tuple[str, ...],
) -> None:
    """Render the input/options panel before the run starts."""
    table = Table.grid(padding=(0, 1))
    table.add_column(style="bold cyan", justify="right")
    table.add_column()
    table.add_row("input:", str(args.input))
    if args.v3:
        table.add_row("pipeline:", "v3 (whole-document, SPECS criteria, gap loop)")
        table.add_row("model:", args.model or "")
        table.add_row("document type:", args.document_type)
        table.add_row("gap-loop cap:", str(args.rounds))
        table.add_row("outputs:", ", ".join(str(o) for o in args.out))
        console.print(Panel(table, title="whetstone v3", border_style="cyan"))
        return
    table.add_row("mode:", args.mode)
    table.add_row("model:", args.model or "[i](deterministic only)[/i]")
    if args.mode == "panel":
        table.add_row("n_experts:", str(args.n_experts))
        if args.panel_disciplines:
            table.add_row("disciplines:", args.panel_disciplines)
    elif args.mode == "guidelines":
        guideline_preview = (
            f"@{args.guidelines[1:]}"
            if args.guidelines.startswith("@")
            else f"<{len(args.guidelines)} chars inline>"
        )
        table.add_row("guidelines:", guideline_preview)
    elif args.mode == "custom":
        criteria_preview = "; ".join(args.criteria) if args.criteria else "(none)"
        table.add_row("criteria:", criteria_preview)
    else:
        table.add_row("perspectives:", ", ".join(perspectives))
        table.add_row("editor:", "on" if args.editor else "off")
        if args.editor:
            table.add_row("editor criteria:", ", ".join(editor_criteria))
        table.add_row("challenge:", "off" if args.no_challenge else "on")
        table.add_row("rounds (cap):", str(args.rounds))
    table.add_row("outputs:", ", ".join(str(o) for o in args.out))
    console.print(Panel(table, title="whetstone v2", border_style="cyan"))


def _print_summary(
    console: Console, result, written: list[Path], *, v3: bool = False
) -> None:
    """Render a Rich summary table after the run completes."""
    m = result.metrics
    table = Table(show_header=False, box=None, padding=(0, 1))
    table.add_column(style="bold cyan", justify="right")
    table.add_column()
    table.add_row("wall time:", f"{m.wall_seconds:.1f}s")
    # llm_calls / reflection rounds are v2 telemetry; v3 doesn't track them yet,
    # so omit rather than print misleading zeros (the [v3.gaps] log shows rounds).
    if not v3:
        table.add_row("LLM calls:", str(m.llm_calls))
    table.add_row("sections:", str(m.sections_processed))
    table.add_row(
        "findings:",
        f"{m.deterministic_findings_count} deterministic + "
        f"{m.investigated_findings_count} investigated "
        f"({m.challenged_findings_count} challenged)",
    )
    table.add_row("edits:", str(m.edits_count))
    table.add_row("author questions:", str(len(result.author_questions)))
    if not v3:
        table.add_row("rounds used:", str(m.reflection_rounds_used))
    table.add_row("outputs:", "\n".join(str(p) for p in written))
    console.print(Panel(table, title="review complete", border_style="green"))


async def _ensure_docx_source(input_arg: str) -> tuple[Path, bool]:
    """Return ``(path, is_temp)`` for a .docx suitable as track-changes baseline.

    If the input is already a local .docx, return its Path with
    ``is_temp=False``. Otherwise harvest the input into markdown and
    write a clean baseline .docx via python-docx; return that path with
    ``is_temp=True`` so the caller knows to clean it up.
    """
    input_path = Path(input_arg) if not _is_url(input_arg) else None
    if input_path is not None and input_path.suffix.lower() == ".docx":
        return input_path, False

    logger = logging.getLogger(_LOGGER_NAME)
    logger.info("[baseline] input is not .docx — generating from harvested text")

    from andamentum.harvest import extract as harvest_extract

    markdown = await harvest_extract(input_arg)
    return _markdown_to_baseline_docx(markdown), True


def _is_url(s: str) -> bool:
    return s.startswith(("http://", "https://", "file://"))


def _markdown_to_baseline_docx(markdown: str) -> Path:
    """Generate a clean baseline .docx from markdown text.

    The output is intentionally minimal — just headings and paragraphs —
    because the goal is to give the track-changes machinery something to
    write into. Anchor matching against this plain prose works because
    ``find_anchor`` is whitespace-tolerant.
    """
    import tempfile

    from docx import Document  # python-docx

    doc = Document()
    for block in markdown.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("#"):
            level = 0
            while level < 9 and level < len(block) and block[level] == "#":
                level += 1
            title = block[level:].strip()
            doc.add_heading(title, level=min(level, 9))
        else:
            doc.add_paragraph(block)

    fd, path = tempfile.mkstemp(suffix=".docx", prefix="whetstone-baseline-")
    import os

    os.close(fd)
    out = Path(path)
    doc.save(str(out))
    return out


# ── Subcommand front-end (Phase F) ──────────────────────────────────────────
#
# The CLI is internally a single flat argparse parser (the canonical surface
# tests build against), with a thin subcommand alias layer in front of it.
# Recognised verbs:
#
#   review         — default; the bare positional invocation also routes here
#                    (`andamentum-whetstone draft.md --model X --out r.md`)
#   panel          — equivalent to `--mode panel`
#   proofread      — shells out to `andamentum.proofread` (no whetstone review
#                    pipeline at all; the two are intentionally disjoint)
#   apply-patches  — equivalent to `--apply-patches PATH`; INPUT and PATCHES
#                    become positional args in the subcommand form
#
# The flat-flag form keeps working for every existing test + script. Users
# who want the subcommand UX get it. v2's mode flag stays callable until
# Phase I deletes v2.

_KNOWN_SUBCOMMANDS: frozenset[str] = frozenset(
    {"review", "panel", "proofread", "apply-patches"}
)


def _rewrite_subcommand(argv: list[str]) -> list[str]:
    """Translate a subcommand-style invocation into the underlying flat-CLI
    argv. Returns ``argv`` unchanged when no subcommand verb is present
    (back-compat for the existing bare-positional form)."""
    if not argv:
        return argv
    head = argv[0]
    rest = argv[1:]
    if head == "review":
        # `review draft.md --model X --out r.md`  →  `draft.md --model X --out r.md`
        return rest
    if head == "panel":
        # `panel draft.md ...`  →  `draft.md --mode panel ...`
        if not rest or rest[0].startswith("-"):
            return ["--mode", "panel", *rest]
        return [rest[0], "--mode", "panel", *rest[1:]]
    if head == "apply-patches":
        # `apply-patches draft.docx --patches p.json --out r.docx`  →
        # `draft.docx --apply-patches p.json --out r.docx`
        # The subcommand requires the patches as a positional or via --patches.
        if not rest or rest[0].startswith("-"):
            _die(
                1,
                "apply-patches subcommand: positional INPUT required.\n"
                "Usage: andamentum-whetstone apply-patches INPUT.docx "
                "--patches PATCHES.json --out OUTPUT.docx",
            )
        out: list[str] = [rest[0]]
        skip_next = False
        for i, tok in enumerate(rest[1:], start=1):
            if skip_next:
                skip_next = False
                continue
            if tok == "--patches":
                if i + 1 >= len(rest):
                    _die(1, "apply-patches: --patches needs a value.")
                out.extend(["--apply-patches", rest[i + 1]])
                skip_next = True
            else:
                out.append(tok)
        return out
    return argv


def _run_proofread_subcommand(argv: list[str]) -> int:
    """The proofread subcommand is a thin shell over ``andamentum.proofread.cli``.
    Whetstone's idea-review pipeline and the deterministic style check are
    intentionally separate jobs (see the project memory entry on this).
    Passes through stdin / file / URL handling unchanged."""
    from andamentum.proofread.cli import main as proofread_main

    return proofread_main(argv)


def main(argv: Sequence[str] | None = None) -> None:
    """Entry point — invoked by the ``andamentum-whetstone`` script.

    Accepts either the bare-positional form (``andamentum-whetstone draft.md
    --model X --out r.md``) or the subcommand form (``andamentum-whetstone
    review draft.md ...`` / ``... panel ...`` / ``... proofread draft.md`` /
    ``... apply-patches draft.docx --patches p.json --out r.docx``)."""
    raw = list(sys.argv[1:] if argv is None else argv)
    if raw and raw[0] == "proofread":
        # proofread short-circuits the whetstone pipeline entirely
        rc = _run_proofread_subcommand(raw[1:])
        raise SystemExit(rc)

    if raw and raw[0] in _KNOWN_SUBCOMMANDS:
        raw = _rewrite_subcommand(raw)

    parser = _build_parser()
    args = parser.parse_args(raw)
    _validate_args(args)

    console = Console(stderr=True)
    _configure_logging(console, args.verbose)

    try:
        if args.apply_patches is not None:
            _apply_patches_only(args, console)
        else:
            asyncio.run(_run(args, console))
    except KeyboardInterrupt:
        _die(1, "interrupted")


def _apply_patches_only(args: argparse.Namespace, console: Console) -> None:
    """Patch-only path: load JSON patches and apply to INPUT.docx → OUT.docx.

    No LLM, no review pipeline. Reuses ``finalize_reviewed_document`` so
    the output is byte-identical to what the full review pipeline would
    produce given the same patches.
    """
    import json

    from .docx.finalization import finalize_reviewed_document
    from .models import DocumentPatch

    raw = json.loads(args.apply_patches.read_text(encoding="utf-8"))
    if not isinstance(raw, list):
        _die(
            1,
            f"patches file must contain a JSON array, got {type(raw).__name__}.",
        )

    try:
        patches = [DocumentPatch(**p) for p in raw]
    except Exception as exc:
        _die(1, f"invalid patch in {args.apply_patches}: {exc}")

    review_summary = ""
    if args.patch_report is not None:
        review_summary = args.patch_report.read_text(encoding="utf-8")

    output_path: Path = args.out[0]
    output_path.parent.mkdir(parents=True, exist_ok=True)

    from .renderers.docx import DEFAULT_AI_AUTHOR

    resolved_author = (
        args.patch_author if args.patch_author is not None else DEFAULT_AI_AUTHOR
    )
    if args.allow_author_override:
        print(
            f"WARNING: attributing AI-generated patches to '{resolved_author}' "
            f"instead of the default '{DEFAULT_AI_AUTHOR}'. Misrepresenting "
            f"AI work as human authorship may constitute research misconduct.",
            file=sys.stderr,
        )

    try:
        _, patch_result = finalize_reviewed_document(
            original_file_path=Path(args.input),
            patches=patches,
            review_summary=review_summary,
            issues_count=len(patches),
            output_path=output_path,
            author=resolved_author,
            use_patch_authors=False,
        )
        # Apply invisible AI-provenance metadata so the modified manuscript
        # carries discoverable AI authorship even without a visible banner.
        from ._watermark import stamp_docx_core_properties

        stamp_docx_core_properties(output_path, model=None)
    except Exception as exc:
        _die(3, f"patch application failed: {exc}")

    _print_apply_patches_summary(console, patch_result, output_path)


def _print_apply_patches_summary(console: Console, result, output_path: Path) -> None:
    """Render a Rich summary table after a patch-only run."""
    table = Table(show_header=False, box=None, padding=(0, 1))
    table.add_column(style="bold cyan", justify="right")
    table.add_column()
    table.add_row("total patches:", str(result.total_patches))
    table.add_row(
        "applied:",
        f"{result.applied_patches} "
        f"({result.applied_edits} edits, {result.applied_comments} comments)",
    )
    table.add_row("failed:", str(len(result.failed_patches)))
    table.add_row("success rate:", f"{result.success_rate:.1f}%")
    table.add_row("output:", str(output_path))
    border = "yellow" if result.failed_patches else "green"
    console.print(Panel(table, title="patches applied", border_style=border))

    if result.failed_patches:
        console.print(
            f"[yellow]warning:[/yellow] {len(result.failed_patches)} "
            f"patch(es) could not be applied:"
        )
        for p in result.failed_patches:
            label = (p.text_pattern or p.patch_type)[:60]
            console.print(f"  • [{p.patch_type}] {label}")


def _configure_logging(console: Console, verbose: bool) -> None:
    """Install a RichHandler on the v2 logger.

    INFO when --verbose, WARNING otherwise so warnings/errors still surface
    even on quiet runs. We attach to our own namespace logger (not root)
    so we don't accidentally hijack pydantic-ai/httpx noise.
    """
    logger = logging.getLogger(_LOGGER_NAME)
    # Avoid duplicate handlers if main() is called twice in one process.
    for h in list(logger.handlers):
        logger.removeHandler(h)
    logger.setLevel(logging.INFO if verbose else logging.WARNING)
    handler = RichHandler(
        console=console,
        show_time=verbose,
        show_path=False,
        show_level=verbose,
        markup=False,
        rich_tracebacks=True,
    )
    handler.setLevel(logging.INFO if verbose else logging.WARNING)
    logger.addHandler(handler)
    logger.propagate = False


if __name__ == "__main__":  # pragma: no cover
    main()
