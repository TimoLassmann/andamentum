"""CLI for whetstone v2 — ``andamentum-whetstone``.

Single positional argument (the input document, anything ``harvest`` can
read), one or more ``--out`` files (format inferred from the extension),
and a small set of well-named options. ``--help`` is exhaustive on
purpose: this is the surface most users will touch.

Exit codes:
    0 — success
    1 — configuration / argument error
    2 — input could not be loaded (harvest failure, file not found)
    3 — review pipeline failed
"""

from __future__ import annotations

import argparse
import asyncio
import logging
import sys
from pathlib import Path
from typing import NoReturn, Sequence

from rich.console import Console
from rich.logging import RichHandler
from rich.panel import Panel
from rich.table import Table


_LOGGER_NAME = "andamentum.whetstone.v2"


_HELP_DESCRIPTION = """\
Sharpen your draft with structured review.

INPUT
    Path or URL to a document. Supported: .pdf, .docx, .html, .md, .txt
    (and any URL; HTML article-vs-listing is auto-detected).

OUTPUTS
    --out FILE         An output file. Repeat for multiple formats.
                       Format inferred from the extension:
                         .md   → markdown report
                         .html → HTML report (via typeset)
                         .docx → Word with track-changes + comments
                       For .docx: if INPUT is .docx we use it directly;
                       otherwise a clean .docx is generated from the
                       harvested text and patches are applied to that.

REVIEW OPTIONS
    --mode {review,panel,guidelines,custom}
                               Pipeline to run.
                                 review (default) — lens-based critical
                                   review with reflection loop.
                                 panel — simulate 3-5 fictional expert
                                   reviewers, scored, with a panel
                                   synthesis. Costs ~2N+2 LLM calls.
                                 guidelines — extract checkable items
                                   from a journal's author guidelines
                                   and evaluate each against the
                                   manuscript. Requires --guidelines.
                                 custom — evaluate the manuscript
                                   against caller-supplied criteria
                                   (one LLM call). Requires --criteria.
    --model MODEL              pydantic-ai model id, e.g.
                               openai:gpt-5.4-nano,
                               ollama:gemma4:31b-nvfp4.
                               Required unless --no-llm is given.
    --editor                   Also generate concrete edits (rewrites).
                               Default off — adds one LLM call per section.
                               Ignored in panel mode.
    --editor-criteria LIST     Comma-separated. Default:
                               clarity,concision,grammar
    --no-challenge             Skip the refutation phase
                               (faster, slightly less reliable findings).
                               Ignored in panel mode.
    --perspectives LIST        Comma-separated lens names.
                               Default: rigorous
                               Available: rigorous, writer, methodology,
                                          statistician
                               Ignored in panel mode.
    --rounds N                 Hard cap on rounds of the reflection–
                               investigation loop. Default: 3.
                               The loop typically exits earlier when the
                               senior reviewer says "nothing more to do".
                               Ignored in panel mode.
    --n-experts N              In panel mode, how many experts to
                               generate. Default 4.
    --panel-disciplines LIST   In panel mode, an explicit comma-separated
                               list of disciplines (skips keyword
                               extraction).
    --guidelines TEXT          In guidelines mode, the journal author
                               guidelines as text. Use ``@path/to/file``
                               to read from a file. Required when
                               ``--mode guidelines``.
    --criteria LIST            In custom mode, semicolon-separated list
                               of criteria (e.g. "originality; depth of
                               literature; clarity of methods"). May be
                               repeated to supply criteria one at a time.
                               Required when ``--mode custom``.
    --no-llm                   Run only the deterministic structural pass
                               (citations, terms, numerics, cross-refs).
                               No --model required. Free, instant.
    -v, --verbose              Print phase-by-phase progress to stderr.
"""


_HELP_EXAMPLES = """\
EXAMPLES
    # Quick deterministic-only check (no LLM, no cost)
    andamentum-whetstone paper.pdf --no-llm --out review.md

    # Full critical review of a PDF, output in three formats
    andamentum-whetstone paper.pdf \\
        --model openai:gpt-5.4-nano \\
        --out review.md --out review.html --out paper.reviewed.docx

    # Edit-mode: get concrete rewrites in a tracked-change Word doc
    andamentum-whetstone draft.docx \\
        --model openai:gpt-5.4-nano --editor \\
        --out draft.reviewed.docx

    # Multi-lens panel on an arXiv paper
    andamentum-whetstone https://arxiv.org/pdf/1901.01753 \\
        --model openai:gpt-5.4-nano \\
        --perspectives rigorous,statistician,writer \\
        --out panel-review.html

    # Local model, deeper reflection loop (4 rounds), all outputs
    andamentum-whetstone manuscript.pdf \\
        --model ollama:gemma4:31b-nvfp4 --rounds 4 --editor \\
        --out review.md --out review.html --out manuscript.reviewed.docx

EXIT CODES
    0  success
    1  configuration / argument error
    2  input could not be loaded (harvest failure, file not found)
    3  review pipeline failed
"""


_SUPPORTED_OUT_EXTENSIONS = {".md", ".html", ".docx"}


def _build_parser() -> argparse.ArgumentParser:
    """Build the argparse parser with the full help text."""
    parser = argparse.ArgumentParser(
        prog="andamentum-whetstone",
        description=_HELP_DESCRIPTION,
        epilog=_HELP_EXAMPLES,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument(
        "input",
        help="Path or URL to the document to review.",
    )
    parser.add_argument(
        "--out",
        action="append",
        type=Path,
        required=True,
        metavar="FILE",
        help="Output file. Repeat for multiple formats. "
        "Format inferred from extension (.md / .html / .docx).",
    )
    parser.add_argument(
        "--mode",
        choices=("review", "panel", "guidelines", "custom"),
        default="review",
        help=(
            "Pipeline to run. review (default): lens-based critical "
            "review. panel: simulate 3-5 fictional expert reviewers. "
            "guidelines: evaluate document against journal author "
            "guidelines (requires --guidelines). custom: evaluate "
            "against caller-supplied criteria (requires --criteria)."
        ),
    )
    parser.add_argument(
        "--model",
        metavar="MODEL",
        help="pydantic-ai model id (e.g. openai:gpt-5.4-nano, "
        "ollama:gemma4:31b-nvfp4). Required unless --no-llm.",
    )
    parser.add_argument(
        "--n-experts",
        type=int,
        default=4,
        metavar="N",
        help="In panel mode, how many experts to generate. Default: 4.",
    )
    parser.add_argument(
        "--panel-disciplines",
        default="",
        metavar="LIST",
        help=(
            "In panel mode, an explicit comma-separated list of "
            "disciplines (skips keyword extraction)."
        ),
    )
    parser.add_argument(
        "--guidelines",
        default="",
        metavar="TEXT",
        help=(
            "In guidelines mode, the journal author guidelines text. "
            "Use ``@path/to/file`` to read from a file. Required when "
            "--mode guidelines."
        ),
    )
    parser.add_argument(
        "--criteria",
        action="append",
        default=None,
        metavar="LIST",
        help=(
            "In custom mode, criteria to evaluate against. Either a "
            "single semicolon-separated string or repeat the flag to "
            "supply criteria one at a time. Required when --mode custom."
        ),
    )
    parser.add_argument(
        "--editor",
        action="store_true",
        help="Also generate concrete edits (rewrites). Adds one LLM call per section.",
    )
    parser.add_argument(
        "--editor-criteria",
        default="clarity,concision,grammar",
        metavar="LIST",
        help="Comma-separated editor criteria. Default: clarity,concision,grammar",
    )
    parser.add_argument(
        "--no-challenge",
        action="store_true",
        help="Skip the refutation phase (faster, slightly less reliable findings).",
    )
    parser.add_argument(
        "--perspectives",
        default="rigorous",
        metavar="LIST",
        help=(
            "Comma-separated lens names. "
            "Available: rigorous, writer, methodology, statistician. "
            "Default: rigorous"
        ),
    )
    parser.add_argument(
        "--rounds",
        type=int,
        default=3,
        metavar="N",
        help=(
            "Hard cap on rounds of the reflection-investigation loop. "
            "Default: 3. The loop typically exits earlier."
        ),
    )
    parser.add_argument(
        "--no-llm",
        action="store_true",
        help="Run only the deterministic structural pass. No --model required.",
    )
    parser.add_argument(
        "--check-novelty",
        action="store_true",
        help=(
            "Verify the manuscript's novelty claims against the literature "
            "via deep_research. Adds 1 + (3-5)*2 LLM calls + several web "
            "fetches. Disabled by default."
        ),
    )
    parser.add_argument(
        "--novelty-search-depth",
        type=int,
        default=2,
        metavar="N",
        help=(
            "Search depth for --check-novelty. 1=quick, 2=balanced, 3=thorough. "
            "Default: 2."
        ),
    )
    parser.add_argument(
        "-v",
        "--verbose",
        action="store_true",
        help="Print progress to stderr.",
    )
    return parser


def _validate_args(args: argparse.Namespace) -> None:
    """Argparse-time validation that argparse can't express directly.

    Raises ``SystemExit(1)`` with a friendly message on any violation.
    """
    if not args.no_llm and not args.model:
        _die(
            1,
            "--model is required unless you pass --no-llm.\n"
            "Examples:\n"
            "  --model openai:gpt-5.4-nano\n"
            "  --model ollama:gemma4:31b-nvfp4",
        )
    if args.no_llm and args.editor:
        _die(
            1,
            "--editor requires --model (the editor agent is an LLM call). "
            "Drop --editor or --no-llm.",
        )
    if args.no_llm and args.mode == "panel":
        _die(
            1,
            "--mode panel requires --model — every panel-mode phase is "
            "an LLM call. Drop --no-llm or use --mode review.",
        )
    if args.no_llm and args.mode == "guidelines":
        _die(
            1,
            "--mode guidelines requires --model — every checkable item "
            "is an LLM call. Drop --no-llm or use --mode review.",
        )
    if args.no_llm and args.mode == "custom":
        _die(
            1,
            "--mode custom requires --model — the custom reviewer is an "
            "LLM call. Drop --no-llm or use --mode review.",
        )
    if args.mode == "guidelines" and not args.guidelines:
        _die(
            1,
            "--mode guidelines requires --guidelines (text or @file).",
        )
    if args.mode == "custom" and not args.criteria:
        _die(
            1,
            "--mode custom requires --criteria (semicolon-separated or "
            "repeated).",
        )
    if args.mode != "guidelines" and args.guidelines:
        _die(
            1,
            f"--guidelines is only valid with --mode guidelines (got "
            f"--mode {args.mode}).",
        )
    if args.mode != "custom" and args.criteria:
        _die(
            1,
            f"--criteria is only valid with --mode custom (got --mode "
            f"{args.mode}).",
        )
    if args.rounds < 1:
        _die(1, "--rounds must be at least 1.")
    if args.n_experts < 1:
        _die(1, "--n-experts must be at least 1.")
    for out in args.out:
        ext = out.suffix.lower()
        if ext not in _SUPPORTED_OUT_EXTENSIONS:
            _die(
                1,
                f"Unsupported output extension {ext!r} for {out}. "
                f"Supported: {', '.join(sorted(_SUPPORTED_OUT_EXTENSIONS))}",
            )


def _die(code: int, message: str) -> NoReturn:
    """Print to stderr and exit with the given code."""
    print(f"error: {message}", file=sys.stderr)
    sys.exit(code)


def _resolve_guidelines(value: str) -> str:
    """Resolve --guidelines value, expanding ``@path`` to file contents.

    A leading ``@`` selects file mode; everything else is treated as a
    literal guidelines string.

    Raises
    ------
    SystemExit
        With exit code 2 if the file cannot be read.
    """
    if value.startswith("@"):
        path = Path(value[1:])
        try:
            return path.read_text(encoding="utf-8")
        except OSError as exc:
            _die(2, f"could not read guidelines file {path}: {exc}")
    return value


def _parse_criteria(values: list[str]) -> list[str]:
    """Parse ``--criteria`` into a flat list of criterion strings.

    Each ``--criteria`` invocation may be a semicolon-separated string;
    multiple invocations stack. Empty / whitespace-only entries are
    silently dropped.
    """
    out: list[str] = []
    for raw in values:
        for chunk in raw.split(";"):
            cleaned = chunk.strip()
            if cleaned:
                out.append(cleaned)
    return out


async def _run(args: argparse.Namespace, console: Console) -> None:
    """Execute the review and write all requested outputs."""
    from .api import review_document
    from .renderers import render_docx, render_html, render_markdown

    perspectives = tuple(p.strip() for p in args.perspectives.split(",") if p.strip())
    editor_criteria = tuple(
        c.strip() for c in args.editor_criteria.split(",") if c.strip()
    )
    panel_disciplines = tuple(
        d.strip() for d in args.panel_disciplines.split(",") if d.strip()
    )
    guidelines_text = _resolve_guidelines(args.guidelines) if args.guidelines else ""
    custom_criteria = _parse_criteria(args.criteria) if args.criteria else None

    logger = logging.getLogger(_LOGGER_NAME)
    if args.verbose:
        _log_run_config(console, args, perspectives, editor_criteria)

    try:
        result = await review_document(
            args.input,
            model=None if args.no_llm else args.model,
            perspectives=perspectives,
            reflection_round_cap=args.rounds,
            challenge=not args.no_challenge,
            editor=args.editor,
            editor_criteria=editor_criteria,
            mode=args.mode,
            n_experts=args.n_experts,
            panel_disciplines=panel_disciplines or None,
            guidelines=guidelines_text,
            custom_criteria=custom_criteria,
            check_novelty=args.check_novelty,
            novelty_search_depth=args.novelty_search_depth,
        )
    except Exception as exc:
        # Distinguish "couldn't load the input" from "review crashed"
        # so users know whether to fix their input or their config.
        msg = str(exc)
        if any(t in msg.lower() for t in ("fetch", "harvest", "no such file", "not found")):
            _die(2, f"could not load input {args.input!r}: {exc}")
        else:
            _die(3, f"review pipeline failed: {exc}")

    # If any output is .docx, prepare a baseline .docx ONCE upfront
    # (rather than per-output) so we don't re-harvest if the user asked
    # for multiple .docx outputs.
    baseline_docx: Path | None = None
    baseline_is_temp = False
    if any(o.suffix.lower() == ".docx" for o in args.out):
        baseline_docx, baseline_is_temp = await _ensure_docx_source(args.input)

    written: list[Path] = []
    try:
        # ── Write each requested output ──────────────────────────────
        for out in args.out:
            ext = out.suffix.lower()
            out.parent.mkdir(parents=True, exist_ok=True)
            if ext == ".md":
                render_markdown(result, out)
            elif ext == ".html":
                render_html(result, out)
            elif ext == ".docx":
                assert baseline_docx is not None  # validated above
                render_docx(result, source_docx_path=baseline_docx, output_path=out)
            written.append(out)
            logger.info("[output] wrote %s", out)
    finally:
        # Clean up the auto-generated baseline docx (if we created one).
        if baseline_is_temp and baseline_docx is not None:
            try:
                baseline_docx.unlink()
            except OSError:
                pass

    _print_summary(console, result, written)


def _log_run_config(
    console: Console,
    args: argparse.Namespace,
    perspectives: tuple[str, ...],
    editor_criteria: tuple[str, ...],
) -> None:
    """Render the input/options panel before the run starts."""
    table = Table.grid(padding=(0, 1))
    table.add_column(style="bold cyan", justify="right")
    table.add_column()
    table.add_row("input:", str(args.input))
    table.add_row("mode:", args.mode)
    table.add_row("model:", args.model or "[i](deterministic only)[/i]")
    if args.mode == "panel":
        table.add_row("n_experts:", str(args.n_experts))
        if args.panel_disciplines:
            table.add_row("disciplines:", args.panel_disciplines)
    elif args.mode == "guidelines":
        guideline_preview = (
            f"@{args.guidelines[1:]}"
            if args.guidelines.startswith("@")
            else f"<{len(args.guidelines)} chars inline>"
        )
        table.add_row("guidelines:", guideline_preview)
    elif args.mode == "custom":
        criteria_preview = "; ".join(args.criteria) if args.criteria else "(none)"
        table.add_row("criteria:", criteria_preview)
    else:
        table.add_row("perspectives:", ", ".join(perspectives))
        table.add_row("editor:", "on" if args.editor else "off")
        if args.editor:
            table.add_row("editor criteria:", ", ".join(editor_criteria))
        table.add_row("challenge:", "off" if args.no_challenge else "on")
        table.add_row("rounds (cap):", str(args.rounds))
    table.add_row("outputs:", ", ".join(str(o) for o in args.out))
    console.print(Panel(table, title="whetstone v2", border_style="cyan"))


def _print_summary(
    console: Console, result, written: list[Path]
) -> None:
    """Render a Rich summary table after the run completes."""
    m = result.metrics
    table = Table(show_header=False, box=None, padding=(0, 1))
    table.add_column(style="bold cyan", justify="right")
    table.add_column()
    table.add_row("wall time:", f"{m.wall_seconds:.1f}s")
    table.add_row("LLM calls:", str(m.llm_calls))
    table.add_row("sections:", str(m.sections_processed))
    table.add_row(
        "findings:",
        f"{m.deterministic_findings_count} deterministic + "
        f"{m.investigated_findings_count} investigated "
        f"({m.challenged_findings_count} challenged)",
    )
    table.add_row("edits:", str(m.edits_count))
    table.add_row("author questions:", str(len(result.author_questions)))
    table.add_row("rounds used:", str(m.reflection_rounds_used))
    table.add_row("outputs:", "\n".join(str(p) for p in written))
    console.print(Panel(table, title="review complete", border_style="green"))


async def _ensure_docx_source(input_arg: str) -> tuple[Path, bool]:
    """Return ``(path, is_temp)`` for a .docx suitable as track-changes baseline.

    If the input is already a local .docx, return its Path with
    ``is_temp=False``. Otherwise harvest the input into markdown and
    write a clean baseline .docx via python-docx; return that path with
    ``is_temp=True`` so the caller knows to clean it up.
    """
    input_path = Path(input_arg) if not _is_url(input_arg) else None
    if input_path is not None and input_path.suffix.lower() == ".docx":
        return input_path, False

    logger = logging.getLogger(_LOGGER_NAME)
    logger.info("[baseline] input is not .docx — generating from harvested text")

    from andamentum.harvest import extract as harvest_extract

    markdown = await harvest_extract(input_arg)
    return _markdown_to_baseline_docx(markdown), True


def _is_url(s: str) -> bool:
    return s.startswith(("http://", "https://", "file://"))


def _markdown_to_baseline_docx(markdown: str) -> Path:
    """Generate a clean baseline .docx from markdown text.

    The output is intentionally minimal — just headings and paragraphs —
    because the goal is to give the track-changes machinery something to
    write into. Anchor matching against this plain prose works because
    ``find_anchor`` is whitespace-tolerant.
    """
    import tempfile

    from docx import Document  # python-docx

    doc = Document()
    for block in markdown.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("#"):
            level = 0
            while level < 9 and level < len(block) and block[level] == "#":
                level += 1
            title = block[level:].strip()
            doc.add_heading(title, level=min(level, 9))
        else:
            doc.add_paragraph(block)

    fd, path = tempfile.mkstemp(suffix=".docx", prefix="whetstone-baseline-")
    import os

    os.close(fd)
    out = Path(path)
    doc.save(str(out))
    return out


def main(argv: Sequence[str] | None = None) -> None:
    """Entry point — invoked by the ``andamentum-whetstone`` script."""
    parser = _build_parser()
    args = parser.parse_args(argv)
    _validate_args(args)

    console = Console(stderr=True)
    _configure_logging(console, args.verbose)

    try:
        asyncio.run(_run(args, console))
    except KeyboardInterrupt:
        _die(1, "interrupted")


def _configure_logging(console: Console, verbose: bool) -> None:
    """Install a RichHandler on the v2 logger.

    INFO when --verbose, WARNING otherwise so warnings/errors still surface
    even on quiet runs. We attach to our own namespace logger (not root)
    so we don't accidentally hijack pydantic-ai/httpx noise.
    """
    logger = logging.getLogger(_LOGGER_NAME)
    # Avoid duplicate handlers if main() is called twice in one process.
    for h in list(logger.handlers):
        logger.removeHandler(h)
    logger.setLevel(logging.INFO if verbose else logging.WARNING)
    handler = RichHandler(
        console=console,
        show_time=verbose,
        show_path=False,
        show_level=verbose,
        markup=False,
        rich_tracebacks=True,
    )
    handler.setLevel(logging.INFO if verbose else logging.WARNING)
    logger.addHandler(handler)
    logger.propagate = False


if __name__ == "__main__":  # pragma: no cover
    main()
