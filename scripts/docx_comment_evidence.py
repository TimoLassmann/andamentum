"""Evidence-collection harness for Word comment rendering — NOT production code.

Modern Microsoft Word (365 / Mac) is not displaying comments that
whetstone writes, even though the comments are structurally valid
classic OOXML (correct `w:` prefix, matched anchors, declared
relationship + content-type). The hypothesis is that modern Word
requires the "modern comment" parts (`commentsExtended.xml`,
`commentsIds.xml`, `people.xml`) and `w14:paraId` linkage, which our
machinery doesn't write.

Rather than guess in production, this script generates a MATRIX of
minimal .docx files that add comments in progressively-more-modern
ways. Open each in Word and record which ones actually show comments.
That gives an empirical truth table of what Word needs.

Variants produced (all built from ONE base doc + 3 comments via the
real whetstone finalisation path, then progressively augmented):

  v1_classic        — exactly what whetstone writes today
                      (comments.xml w: prefix, anchors, rel, content-type)
  v2_paraid         — v1 + w14:paraId/w14:textId on each comment paragraph
  v3_extended       — v2 + word/commentsExtended.xml (w15) + rel + content-type
  v4_modern         — v3 + word/commentsIds.xml (w16cid) + word/people.xml

Run:
    uv run python scripts/docx_comment_evidence.py
Output:
    /tmp/docx_evidence/v1_classic.docx ... v4_modern.docx

Then open each in Word for Mac (Review → All Markup) and note which
show the 3 comments. Report the results back.
"""

from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

OUT_DIR = Path("/tmp/docx_evidence")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"
W16CID_NS = "http://schemas.microsoft.com/office/word/2016/wordml/cid"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"

REL_COMMENTS_EXT = "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"
REL_COMMENTS_IDS = "http://schemas.microsoft.com/office/2016/09/relationships/commentsIds"
REL_PEOPLE = "http://schemas.microsoft.com/office/2011/relationships/people"

CT_COMMENTS_EXT = "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml"
CT_COMMENTS_IDS = "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsIds+xml"
CT_PEOPLE = "application/vnd.openxmlformats-officedocument.wordprocessingml.people+xml"


# ---------------------------------------------------------------------------
# Base document + classic comments (via the real whetstone machinery)
# ---------------------------------------------------------------------------


def build_v1_classic(out: Path) -> list[str]:
    """Build the base doc with 3 comments using whetstone's real finalisation
    path. Returns the comment IDs present in the output."""
    from docx import Document as DocxDocument

    from andamentum.whetstone.docx.finalization import finalize_reviewed_document
    from andamentum.whetstone.models import DocumentPatch

    src = OUT_DIR / "_base_source.docx"
    doc = DocxDocument()
    doc.add_paragraph(
        "The system uses an iterative search strategy with "
        "verified quotes and entity pairs."
    )
    doc.add_paragraph(
        "We evaluated recall against a gold standard and tested resilience "
        "to irrelevant information."
    )
    doc.add_paragraph(
        "The discussion overstates novelty without a concrete baseline "
        "comparison."
    )
    doc.save(str(src))

    patches = [
        DocumentPatch(
            patch_type="comment",
            text_pattern="iterative search strategy",
            comment_text="Comment ONE: clarify what 'iterative' means here.",
            explanation="x",
        ),
        DocumentPatch(
            patch_type="comment",
            text_pattern="gold standard",
            comment_text="Comment TWO: define how the gold standard was built.",
            explanation="y",
        ),
        DocumentPatch(
            patch_type="comment",
            text_pattern="overstates novelty",
            comment_text="Comment THREE: soften or support this claim.",
            explanation="z",
        ),
    ]
    finalize_reviewed_document(
        original_file_path=src,
        patches=patches,
        output_path=out,
        author="andamentum-whetstone (AI)",
    )

    # Read back the comment ids from the produced file
    with zipfile.ZipFile(out) as z:
        comments_xml = z.read("word/comments.xml").decode("utf-8")
    cids = re.findall(r'<w:comment w:id="(\d+)"', comments_xml)
    return cids


# ---------------------------------------------------------------------------
# Zip surgery helpers
# ---------------------------------------------------------------------------


def _read_all(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as z:
        return {n: z.read(n) for n in z.namelist()}


def _write_all(path: Path, parts: dict[str, bytes]) -> None:
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in parts.items():
            z.writestr(name, data)


def _para_id(cid: str) -> str:
    """Deterministic 8-hex-digit paraId for a comment id (Word-style)."""
    return f"{int(cid):08X}"


# ---------------------------------------------------------------------------
# v2 — add w14:paraId / w14:textId to each comment's first paragraph
# ---------------------------------------------------------------------------


def build_v2_paraid(v1: Path, out: Path, cids: list[str]) -> None:
    parts = _read_all(v1)
    comments = parts["word/comments.xml"].decode("utf-8")

    # Declare w14 on the comments root.
    comments = comments.replace(
        f'<w:comments xmlns:w="{W_NS}">',
        f'<w:comments xmlns:w="{W_NS}" xmlns:w14="{W14_NS}">',
    )

    # For each comment, stamp w14:paraId/textId onto the FIRST <w:p> after the
    # comment's opening tag. We rewrite each comment block's first <w:p>.
    def stamp(match: re.Match) -> str:
        cid = match.group("cid")
        head = match.group(0)
        pid = _para_id(cid)
        # add attributes to the first <w:p> (which may be "<w:p>" exactly)
        return head.replace(
            "<w:p>",
            f'<w:p w14:paraId="{pid}" w14:textId="{pid}">',
            1,
        )

    # Match each whole comment element so the .replace only affects its first <w:p>.
    comments = re.sub(
        r'<w:comment w:id="(?P<cid>\d+)".*?</w:comment>',
        stamp,
        comments,
        flags=re.DOTALL,
    )
    parts["word/comments.xml"] = comments.encode("utf-8")
    _write_all(out, parts)


# ---------------------------------------------------------------------------
# v3 — add commentsExtended.xml + relationship + content-type
# ---------------------------------------------------------------------------


def build_v3_extended(v2: Path, out: Path, cids: list[str]) -> None:
    parts = _read_all(v2)

    # 1) commentsExtended.xml
    ex_entries = "".join(
        f'<w15:commentEx w15:paraId="{_para_id(c)}" w15:done="0"/>' for c in cids
    )
    commentsExtended = (
        f"<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n"
        f'<w15:commentsEx xmlns:w15="{W15_NS}">{ex_entries}</w15:commentsEx>'
    )
    parts["word/commentsExtended.xml"] = commentsExtended.encode("utf-8")

    # 2) relationship in word/_rels/document.xml.rels
    rels_name = "word/_rels/document.xml.rels"
    rels = parts[rels_name].decode("utf-8")
    if "commentsExtended.xml" not in rels:
        rel = (
            f'<Relationship Id="rIdCommentsExt" Type="{REL_COMMENTS_EXT}" '
            f'Target="commentsExtended.xml"/>'
        )
        rels = rels.replace("</Relationships>", rel + "</Relationships>")
        parts[rels_name] = rels.encode("utf-8")

    # 3) content type override
    ct = parts["[Content_Types].xml"].decode("utf-8")
    if "commentsExtended" not in ct:
        override = (
            f'<Override PartName="/word/commentsExtended.xml" '
            f'ContentType="{CT_COMMENTS_EXT}"/>'
        )
        ct = ct.replace("</Types>", override + "</Types>")
        parts["[Content_Types].xml"] = ct.encode("utf-8")

    _write_all(out, parts)


# ---------------------------------------------------------------------------
# v4 — add commentsIds.xml + people.xml + relationships + content-types
# ---------------------------------------------------------------------------


def build_v4_modern(v3: Path, out: Path, cids: list[str], author: str) -> None:
    parts = _read_all(v3)

    # commentsIds.xml
    id_entries = "".join(
        f'<w16cid:commentId w16cid:paraId="{_para_id(c)}" '
        f'w16cid:durableId="{int(c):08X}"/>'
        for c in cids
    )
    commentsIds = (
        f"<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n"
        f'<w16cid:commentsIds xmlns:w16cid="{W16CID_NS}">{id_entries}</w16cid:commentsIds>'
    )
    parts["word/commentsIds.xml"] = commentsIds.encode("utf-8")

    # people.xml
    people = (
        f"<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n"
        f'<w15:people xmlns:w15="{W15_NS}" xmlns:w="{W_NS}">'
        f'<w15:person w15:author="{author}">'
        f'<w15:presenceInfo w15:providerId="None" w15:userId="{author}"/>'
        f"</w15:person></w15:people>"
    )
    parts["word/people.xml"] = people.encode("utf-8")

    # relationships
    rels_name = "word/_rels/document.xml.rels"
    rels = parts[rels_name].decode("utf-8")
    additions = ""
    if "commentsIds.xml" not in rels:
        additions += (
            f'<Relationship Id="rIdCommentsIds" Type="{REL_COMMENTS_IDS}" '
            f'Target="commentsIds.xml"/>'
        )
    if "people.xml" not in rels:
        additions += (
            f'<Relationship Id="rIdPeople" Type="{REL_PEOPLE}" '
            f'Target="people.xml"/>'
        )
    if additions:
        rels = rels.replace("</Relationships>", additions + "</Relationships>")
        parts[rels_name] = rels.encode("utf-8")

    # content types
    ct = parts["[Content_Types].xml"].decode("utf-8")
    ct_add = ""
    if "commentsIds" not in ct:
        ct_add += f'<Override PartName="/word/commentsIds.xml" ContentType="{CT_COMMENTS_IDS}"/>'
    if "people" not in ct:
        ct_add += f'<Override PartName="/word/people.xml" ContentType="{CT_PEOPLE}"/>'
    if ct_add:
        ct = ct.replace("</Types>", ct_add + "</Types>")
        parts["[Content_Types].xml"] = ct.encode("utf-8")

    _write_all(out, parts)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> None:
    if OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    OUT_DIR.mkdir(parents=True)

    author = "andamentum-whetstone (AI)"

    v1 = OUT_DIR / "v1_classic.docx"
    v2 = OUT_DIR / "v2_paraid.docx"
    v3 = OUT_DIR / "v3_extended.docx"
    v4 = OUT_DIR / "v4_modern.docx"

    print("Building v1_classic (whetstone's current output)...")
    cids = build_v1_classic(v1)
    print(f"  comment ids: {cids}")

    print("Building v2_paraid (+ w14:paraId)...")
    build_v2_paraid(v1, v2, cids)

    print("Building v3_extended (+ commentsExtended.xml)...")
    build_v3_extended(v2, v3, cids)

    print("Building v4_modern (+ commentsIds.xml + people.xml)...")
    build_v4_modern(v3, v4, cids, author)

    print()
    print("Done. Open each in Word for Mac (Review tab → All Markup) and note")
    print("which show all 3 comments:")
    for f in (v1, v2, v3, v4):
        print(f"  {f}")
    print()
    print("Each file contains exactly 3 comments anchored to:")
    print("  1. 'iterative search strategy'")
    print("  2. 'gold standard'")
    print("  3. 'overstates novelty'")


if __name__ == "__main__":
    main()
