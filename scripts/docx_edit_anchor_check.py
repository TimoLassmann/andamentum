"""Verify that track-change (text_edit) anchoring is actually precise.

Before modelling the comment-anchoring fix on the edit path, we must
confirm the edit path really does land tracked changes on the exact
target text (and not, say, on the whole paragraph). Tests both a
single-run paragraph and a multi-run paragraph (mimicking LibreOffice,
which splits text across many <w:r> runs — the structure of the real
manuscript).

Run:
    uv run python scripts/docx_edit_anchor_check.py
"""

from __future__ import annotations

import re
import zipfile
from pathlib import Path

OUT = Path("/tmp/docx_edit_check")


def _ins_del_spans(docx_path: Path) -> tuple[list[str], list[str]]:
    """Return (deleted_texts, inserted_texts) from the document body."""
    with zipfile.ZipFile(docx_path) as z:
        doc = z.read("word/document.xml").decode("utf-8")
    # <w:del>...<w:delText>X</w:delText>...</w:del>
    dels = re.findall(r"<w:del\b.*?</w:del>", doc, re.DOTALL)
    deleted = ["".join(re.findall(r"<w:delText[^>]*>(.*?)</w:delText>", d, re.DOTALL)) for d in dels]
    inss = re.findall(r"<w:ins\b.*?</w:ins>", doc, re.DOTALL)
    inserted = ["".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", s, re.DOTALL)) for s in inss]
    return deleted, inserted


def case_single_run() -> None:
    from docx import Document as DocxDocument
    from andamentum.whetstone.docx.finalization import finalize_reviewed_document
    from andamentum.whetstone.models import DocumentPatch

    src = OUT / "single_src.docx"
    out = OUT / "single_out.docx"
    doc = DocxDocument()
    doc.add_paragraph(
        "The methods were significantly robust and the gold standard was "
        "carefully built from expert annotations."
    )
    doc.save(str(src))

    patch = DocumentPatch(
        patch_type="text_edit",
        text_pattern="gold standard",
        new_text="reference set",
        explanation="terminology",
    )
    _, res = finalize_reviewed_document(
        original_file_path=src, patches=[patch], output_path=out, author="test"
    )
    deleted, inserted = _ins_del_spans(out)
    print("── CASE 1: single-run paragraph ──")
    print(f"  applied_patches: {res.applied_patches}/{res.total_patches}")
    print(f"  deleted spans:  {deleted}")
    print(f"  inserted spans: {inserted}")
    ok = any("gold standard" in d for d in deleted) and any("reference set" in i for i in inserted)
    # precise = ONLY the target was deleted, not the whole paragraph
    precise = deleted == ["gold standard"] and inserted == ["reference set"]
    print(f"  VERDICT: anchored={ok}  precise(only-target)={precise}")
    print()


def case_multi_run() -> None:
    """Simulate a LibreOffice-style paragraph: target text split across runs."""
    from docx import Document as DocxDocument
    from andamentum.whetstone.docx.finalization import finalize_reviewed_document
    from andamentum.whetstone.models import DocumentPatch

    src = OUT / "multi_src.docx"
    out = OUT / "multi_out.docx"
    doc = DocxDocument()
    p = doc.add_paragraph()
    # Each add_run is a separate <w:r>; "gold standard" is split across two runs.
    for chunk in ["The methods were ", "significantly ", "robust and the ", "gold ", "standard ", "was carefully built."]:
        p.add_run(chunk)
    doc.save(str(src))

    patch = DocumentPatch(
        patch_type="text_edit",
        text_pattern="gold standard",
        new_text="reference set",
        explanation="terminology",
    )
    _, res = finalize_reviewed_document(
        original_file_path=src, patches=[patch], output_path=out, author="test"
    )
    deleted, inserted = _ins_del_spans(out)
    print("── CASE 2: multi-run paragraph (LibreOffice-style, target split across runs) ──")
    print(f"  applied_patches: {res.applied_patches}/{res.total_patches}")
    print(f"  deleted spans:  {deleted}")
    print(f"  inserted spans: {inserted}")
    ok = any("gold" in d for d in deleted)
    precise = deleted == ["gold standard"] and inserted == ["reference set"]
    print(f"  VERDICT: anchored={ok}  precise(only-target)={precise}")
    print()


def case_pattern_not_found() -> None:
    from docx import Document as DocxDocument
    from andamentum.whetstone.docx.finalization import finalize_reviewed_document
    from andamentum.whetstone.models import DocumentPatch

    src = OUT / "miss_src.docx"
    out = OUT / "miss_out.docx"
    doc = DocxDocument()
    doc.add_paragraph("The methods were robust and the dataset was carefully built.")
    doc.save(str(src))

    patch = DocumentPatch(
        patch_type="text_edit",
        text_pattern="gold standard",  # NOT present
        new_text="reference set",
        explanation="terminology",
    )
    _, res = finalize_reviewed_document(
        original_file_path=src, patches=[patch], output_path=out, author="test"
    )
    deleted, inserted = _ins_del_spans(out)
    print("── CASE 3: text_pattern NOT in document ──")
    print(f"  applied_patches: {res.applied_patches}/{res.total_patches}")
    print(f"  failed_patches: {len(res.failed_patches)}")
    print(f"  deleted spans:  {deleted}")
    print(f"  inserted spans: {inserted}")
    print()


def main() -> None:
    import shutil
    if OUT.exists():
        shutil.rmtree(OUT)
    OUT.mkdir(parents=True)
    case_single_run()
    case_multi_run()
    case_pattern_not_found()


if __name__ == "__main__":
    main()
